from flask import request, jsonify
import dbconnect
from datetime import datetime
import calendar
import logging

from flask import request, jsonify, send_file
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT


logger = logging.getLogger(__name__)

def register_timesheets_reports_routes(app):

    # =========================================================
    # 0) Business Units dropdown (Company)
    # =========================================================
    @app.route('/getbusinessunitslist', methods=['GET'])
    def getbusinessunitslist():

        if 'org_code' in request.args:
            org_code = request.args['org_code']
        else:
            return "Error: No Organization Code field provided. Please specify it."

        active_only = request.args.get('active_only', '1')

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            sql1 = """
                SELECT
                    code,
                    description,
                    status,
                    org_code
                FROM business_units
                WHERE
                    org_code = %s
                    AND (%s = '0' OR status = 1)
                ORDER BY code
            """
            data1 = (org_code, active_only)

            cur.execute(sql1, data1)
            result = cur.fetchall()

            app.config['JSONIFY_PRETTYPRINT_REGULAR'] = True

            if result:
                return jsonify({"message": "Business units retrieved successfully", "result": result}), 200
            else:
                return jsonify({"message": "No business units found", "result": []}), 404

        except Exception as e:
            return jsonify({"message": "Failed to retrieve business units", "error": str(e)}), 500


    # =========================================================
    # 0b) Services dropdown (Scope)
    # =========================================================
    @app.route('/getserviceslist', methods=['GET'])
    def getserviceslist():

        if 'org_code' in request.args:
            org_code = request.args['org_code']
        else:
            return "Error: No Organization Code field provided. Please specify it."

        active_only = request.args.get('active_only', '1')

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            sql1 = """
                SELECT
                    service_id,
                    description,
                    instruction,
                    status,
                    sequence,
                    org_code
                FROM services
                WHERE
                    org_code = %s
                    AND (%s = '0' OR status = 1)
                ORDER BY sequence, description, service_id
            """
            data1 = (org_code, active_only)

            cur.execute(sql1, data1)
            result = cur.fetchall()

            app.config['JSONIFY_PRETTYPRINT_REGULAR'] = True

            if result:
                return jsonify({"message": "Services retrieved successfully", "result": result}), 200
            else:
                return jsonify({"message": "No services found", "result": []}), 404

        except Exception as e:
            return jsonify({"message": "Failed to retrieve services", "error": str(e)}), 500


    # =========================================================
    # 1) Employees Monthly Summary (main table + KPI)
    # ✅ Employee search now matches firstname/lastname (not email)
    # =========================================================
    @app.route('/timesheets_employees_monthly_summary', methods=['GET'])
    def timesheets_employees_monthly_summary():

        if 'org_code' in request.args:
            org_code = request.args['org_code']
        else:
            return "Error: No Organization Code field provided. Please specify it."

        if 'month' in request.args:
            month = request.args['month']
        else:
            return "Error: No Month field provided. Please specify it. Format: YYYY-MM"

        company = request.args.get('company', '')
        scope = request.args.get('scope', '')
        wo_search = request.args.get('wo_search', '')
        employee_search = request.args.get('employee_search', '')

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            sql1 = """
                WITH base AS (
                    SELECT
                        t.email_address AS employee_email,
                        u.firstname AS firstname,
                        u.lastname AS lastname,
                        wo.business_unit AS company,
                        wo.wo_code AS wo_code,
                        wo.project_name AS project_name,
                        t.start AS start_dt,
                        t.end AS end_dt,
                        t.hours AS hours,
                        rs.service_id AS service_id
                    FROM timesheets t
                    LEFT JOIN app_users u
                        ON t.email_address = u.user
                    AND t.org_code = u.org_code
                    LEFT JOIN work_orders wo
                        ON t.wo_code = wo.wo_code
                    LEFT JOIN requested_services rs
                        ON wo.wr_id = rs.wr_id
                    WHERE
                        t.org_code = %s
                        AND DATE_FORMAT(t.start, '%%Y-%%m') = %s
                        AND (%s = '' OR wo.business_unit = %s)
                        AND (%s = '' OR rs.service_id = %s)
                        AND (
                            %s = '' OR
                            t.wo_code LIKE CONCAT('%%', %s, '%%') OR
                            wo.wo_description LIKE CONCAT('%%', %s, '%%') OR
                            wo.project_name LIKE CONCAT('%%', %s, '%%')
                        )
                        AND (
                            %s = '' OR
                            u.firstname LIKE CONCAT('%%', %s, '%%') OR
                            u.lastname LIKE CONCAT('%%', %s, '%%') OR
                            CONCAT(
                                COALESCE(u.firstname, ''),
                                ' ',
                                COALESCE(u.lastname, '')
                            ) LIKE CONCAT('%%', %s, '%%')
                        )
                ),
                emp_agg AS (
                    SELECT
                        employee_email,
                        MAX(firstname) AS firstname,
                        MAX(lastname) AS lastname,
                        project_name,
                        SUM(hours) AS total_hours,
                        COUNT(DISTINCT DATE(start_dt)) AS workdays_logged,
                        COUNT(DISTINCT wo_code) AS wo_count,
                        MAX(end_dt) AS last_log_date
                    FROM base
                    GROUP BY employee_email
                ),
                scope_rank AS (
                    SELECT
                        employee_email,
                        service_id,
                        SUM(hours) AS scope_hours,
                        ROW_NUMBER() OVER (
                            PARTITION BY employee_email
                            ORDER BY SUM(hours) DESC
                        ) AS rn
                    FROM base
                    GROUP BY employee_email, service_id
                ),
                company_rank AS (
                    SELECT
                        employee_email,
                        company,
                        SUM(hours) AS company_hours,
                        ROW_NUMBER() OVER (
                            PARTITION BY employee_email
                            ORDER BY SUM(hours) DESC
                        ) AS rn
                    FROM base
                    GROUP BY employee_email, company
                )
                SELECT
                    a.employee_email,
                    a.firstname,
                    a.lastname,
                    CONCAT(
                        COALESCE(a.firstname, ''),
                        ' ',
                        COALESCE(a.lastname, '')
                    ) AS full_name,
                    a.project_name,
                    a.total_hours,
                    a.workdays_logged,
                    a.wo_count,
                    COALESCE(s.service_id, '-') AS top_scope,
                    COALESCE(c.company, '-') AS top_company,
                    a.last_log_date
                FROM emp_agg a
                LEFT JOIN scope_rank s
                    ON a.employee_email = s.employee_email AND s.rn = 1
                LEFT JOIN company_rank c
                    ON a.employee_email = c.employee_email AND c.rn = 1
                ORDER BY a.total_hours DESC
            """

            data1 = (
                org_code,
                month,
                company, company,
                scope, scope,
                wo_search, wo_search, wo_search, wo_search,
                # ✅ employee_search reused for firstname/lastname/full name matching
                employee_search, employee_search, employee_search, employee_search
            )

            cur.execute(sql1, data1)
            result = cur.fetchall()

            app.config['JSONIFY_PRETTYPRINT_REGULAR'] = True

            if result:
                return jsonify({
                    "message": "Employees monthly timesheets summary retrieved successfully",
                    "result": result
                }), 200
            else:
                return jsonify({"message": "No records found", "result": []}), 404

        except Exception as e:
            return jsonify({
                "message": "Failed to retrieve employees monthly timesheets summary",
                "error": str(e)
            }), 500


    # =========================================================
    # 2) Employee Monthly Entries (drilldown dialog)
    # =========================================================
    @app.route('/timesheets_employee_monthly_entries', methods=['GET'])
    def timesheets_employee_monthly_entries():

        if 'org_code' in request.args:
            org_code = request.args['org_code']
        else:
            return "Error: No Organization Code field provided. Please specify it."

        if 'month' in request.args:
            month = request.args['month']
        else:
            return "Error: No Month field provided. Please specify it. Format: YYYY-MM"

        if 'email_address' in request.args:
            email_address = request.args['email_address']
        else:
            return "Error: No Email Address field provided. Please specify it."

        company = request.args.get('company', '')
        scope = request.args.get('scope', '')
        wo_search = request.args.get('wo_search', '')

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            sql1 = """
                SELECT
                    t.id AS id,
                    t.work_log_id AS work_log_id,
                    t.email_address AS email_address,
                    u.firstname AS firstname,
                    u.lastname AS lastname,
                    CONCAT(
                        COALESCE(u.firstname, ''),
                        ' ',
                        COALESCE(u.lastname, '')
                    ) AS full_name,
                    t.task AS task,
                    t.start AS start_datetime,
                    t.end AS end_datetime,
                    t.hours AS hours,
                    t.created_datetime AS created_datetime,
                    t.org_code AS org_code,
                    t.task_type AS task_type,
                    t.wo_code AS wo_code,
                    t.activities AS activities,
                    t.remarks AS remarks,
                    wo.business_unit AS company,
                    wo.wo_description AS wo_description,
                    wo.project_name AS project_name,
                    rs.service_id AS service_id
                FROM timesheets t
                LEFT JOIN app_users u
                    ON t.email_address = u.user
                AND t.org_code = u.org_code
                LEFT JOIN work_orders wo
                    ON t.wo_code = wo.wo_code AND t.org_code = wo.org_code
                LEFT JOIN requested_services rs
                    ON wo.wr_id = rs.wr_id AND wo.org_code = rs.org_code
                WHERE
                    t.org_code = %s
                    AND DATE_FORMAT(t.start, '%%Y-%%m') = %s
                    AND t.email_address = %s
                    AND (%s = '' OR wo.business_unit = %s)
                    AND (%s = '' OR rs.service_id = %s)
                    AND (
                        %s = '' OR
                        t.wo_code LIKE CONCAT('%%', %s, '%%') OR
                        wo.wo_description LIKE CONCAT('%%', %s, '%%') OR
                        wo.project_name LIKE CONCAT('%%', %s, '%%')
                    )
                ORDER BY t.start DESC
            """

            data1 = (
                org_code,
                month,
                email_address,
                company, company,
                scope, scope,
                wo_search, wo_search, wo_search, wo_search
            )

            cur.execute(sql1, data1)
            result = cur.fetchall()

            app.config['JSONIFY_PRETTYPRINT_REGULAR'] = True

            if result:
                return jsonify({
                    "message": "Employee monthly timesheet entries retrieved successfully",
                    "result": result
                }), 200
            else:
                return jsonify({"message": "No records found", "result": []}), 404

        except Exception as e:
            return jsonify({
                "message": "Failed to retrieve employee monthly timesheet entries",
                "error": str(e)
            }), 500


    # =========================================================
    # 3) Companies dropdown (distinct companies present in timesheets)
    # =========================================================
    @app.route('/timesheets_companies_list', methods=['GET'])
    def timesheets_companies_list():

        if 'org_code' in request.args:
            org_code = request.args['org_code']
        else:
            return "Error: No Organization Code field provided. Please specify it."

        month = request.args.get('month', '')

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            sql1 = """
                SELECT DISTINCT
                    wo.business_unit AS company
                FROM timesheets t
                LEFT JOIN work_orders wo
                    ON t.wo_code = wo.wo_code AND t.org_code = wo.org_code
                WHERE
                    t.org_code = %s
                    AND (%s = '' OR DATE_FORMAT(t.start, '%%Y-%%m') = %s)
                    AND wo.business_unit IS NOT NULL
                    AND wo.business_unit <> ''
                ORDER BY wo.business_unit
            """

            data1 = (org_code, month, month)

            cur.execute(sql1, data1)
            result = cur.fetchall()

            app.config['JSONIFY_PRETTYPRINT_REGULAR'] = True

            if result:
                return jsonify({"message": "Companies retrieved successfully", "result": result}), 200
            else:
                return jsonify({"message": "No companies found", "result": []}), 404

        except Exception as e:
            return jsonify({"message": "Failed to retrieve companies", "error": str(e)}), 500


    # =========================================================
    # 4) Scopes dropdown (distinct service_id present in timesheets)
    # =========================================================
    @app.route('/timesheets_scopes_list', methods=['GET'])
    def timesheets_scopes_list():

        if 'org_code' in request.args:
            org_code = request.args['org_code']
        else:
            return "Error: No Organization Code field provided. Please specify it."

        month = request.args.get('month', '')

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            sql1 = """
                SELECT DISTINCT
                    rs.service_id AS service_id
                FROM timesheets t
                LEFT JOIN work_orders wo
                    ON t.wo_code = wo.wo_code AND t.org_code = wo.org_code
                LEFT JOIN requested_services rs
                    ON wo.wr_id = rs.wr_id AND wo.org_code = rs.org_code
                WHERE
                    t.org_code = %s
                    AND (%s = '' OR DATE_FORMAT(t.start, '%%Y-%%m') = %s)
                    AND rs.service_id IS NOT NULL
                ORDER BY rs.service_id
            """

            data1 = (org_code, month, month)

            cur.execute(sql1, data1)
            result = cur.fetchall()

            app.config['JSONIFY_PRETTYPRINT_REGULAR'] = True

            if result:
                return jsonify({"message": "Scopes retrieved successfully", "result": result}), 200
            else:
                return jsonify({"message": "No scopes found", "result": []}), 404

        except Exception as e:
            return jsonify({"message": "Failed to retrieve scopes", "error": str(e)}), 500


    # =========================================================
    # 5) Companies dropdown (from business_units table)
    # =========================================================
    @app.route('/timesheets_business_units_dropdown', methods=['GET'])
    def timesheets_business_units_dropdown():

        if 'org_code' in request.args:
            org_code = request.args['org_code']
        else:
            return "Error: No Organization Code field provided. Please specify it."

        active_only = request.args.get('active_only', '1')

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            sql1 = """
                SELECT
                    code,
                    description
                FROM business_units
                WHERE
                    org_code = %s
                    AND (%s = '0' OR status = 1)
                ORDER BY code
            """
            data1 = (org_code, active_only)

            cur.execute(sql1, data1)
            result = cur.fetchall()

            app.config['JSONIFY_PRETTYPRINT_REGULAR'] = True

            if result:
                return jsonify({"message": "Business units dropdown retrieved successfully", "result": result}), 200
            else:
                return jsonify({"message": "No business units found", "result": []}), 404

        except Exception as e:
            return jsonify({"message": "Failed to retrieve business units dropdown", "error": str(e)}), 500


    # =========================================================
    # 6) Scopes dropdown (from services table)
    # =========================================================
    @app.route('/timesheets_services_dropdown', methods=['GET'])
    def timesheets_services_dropdown():

        if 'org_code' in request.args:
            org_code = request.args['org_code']
        else:
            return "Error: No Organization Code field provided. Please specify it."

        active_only = request.args.get('active_only', '1')

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            sql1 = """
                SELECT
                    service_id,
                    description
                FROM services
                WHERE
                    org_code = %s
                    AND (%s = '0' OR status = 1)
                ORDER BY sequence, description, service_id
            """
            data1 = (org_code, active_only)

            cur.execute(sql1, data1)
            result = cur.fetchall()

            app.config['JSONIFY_PRETTYPRINT_REGULAR'] = True

            if result:
                return jsonify({"message": "Services dropdown retrieved successfully", "result": result}), 200
            else:
                return jsonify({"message": "No services found", "result": []}), 404

        except Exception as e:
            return jsonify({"message": "Failed to retrieve services dropdown", "error": str(e)}), 500


    # ---------------------------
    # GET SERVICES MASTER (left pane)
    # ---------------------------
    @app.route('/getservicesmaster', methods=['GET'])
    def getservicesmaster():
        if 'org_code' in request.args:
            org_code = request.args['org_code']
        else:
            return jsonify({"error": "No Organization Code field provided. Please specify it."}), 400

        q = request.args.get('q', '').strip()
        active_only = request.args.get('active_only', '').strip()  # "1" or ""

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            where = "WHERE org_code = %s"
            params = [org_code]

            if q != "":
                where += " AND description LIKE %s"
                params.append("%" + q + "%")

            if active_only == "1":
                where += " AND status = 1"

            sql = f"""
                SELECT
                    service_id,
                    description,
                    instruction,
                    status,
                    sequence,
                    org_code
                FROM services
                {where}
                ORDER BY sequence ASC, description ASC
            """

            cur.execute(sql, tuple(params))
            results = cur.fetchall()

            cur.close()
            conn.close()

            return jsonify({"result": results})

        except Exception as e:
            return jsonify({"error": str(e)}), 500


    # ---------------------------
    # GET SERVICE DETAILS (right)
    # ---------------------------
    @app.route('/getservicedetails', methods=['GET'])
    def getservicedetails():
        if 'service_id' in request.args:
            service_id = request.args['service_id']
        else:
            return jsonify({"error": "No Service ID field provided. Please specify it."}), 400

        if 'org_code' in request.args:
            org_code = request.args['org_code']
        else:
            return jsonify({"error": "No Organization Code field provided. Please specify it."}), 400

        active_only = request.args.get('active_only', '').strip()  # "1" or ""

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            where = "WHERE service_id = %s AND org_code = %s"
            params = [service_id, org_code]

            if active_only == "1":
                where += " AND status = 1"

            sql = f"""
                SELECT
                    detail_id,
                    service_id,
                    description,
                    status,
                    sequence,
                    org_code
                FROM service_details
                {where}
                ORDER BY sequence ASC, description ASC
            """

            cur.execute(sql, tuple(params))
            results = cur.fetchall()

            cur.close()
            conn.close()

            return jsonify({"result": results})

        except Exception as e:
            return jsonify({"error": str(e)}), 500


    # --- get timesheet entries by Work Order (optional email filter) ---
    @app.route('/gettimesheets', methods=['GET'])
    def gettimesheets():
        if 'org_code' not in request.args:
            return jsonify({'error': 'No Organization Code field provided.'}), 400

        if 'wo_code' not in request.args:
            return jsonify({'error': 'No Work Order Code provided.'}), 400

        org_code = request.args['org_code']
        wo_code = request.args['wo_code']
        email_address = request.args.get('email_address', '').strip()

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()  # DictCursor assumed

            sql = """
                SELECT
                    id,
                    email_address,
                    task,
                    start,
                    end,
                    hours,
                    created_datetime,
                    org_code,
                    task_type,
                    wo_code,
                    activities,
                    remarks
                FROM timesheets
                WHERE org_code = %s
                AND wo_code = %s
            """

            params = [org_code, wo_code]

            # ✅ OPTIONAL email filter
            if email_address != "":
                sql += " AND email_address = %s "
                params.append(email_address)

            sql += " ORDER BY start DESC, id DESC "

            cur.execute(sql, tuple(params))
            rows = cur.fetchall()   # already dicts

            return jsonify({'result': rows}), 200

        except Exception as e:
            print("gettimesheets error:", e)
            return jsonify({'error': str(e)}), 500


    @app.route('/gettimesheets_user', methods=['GET'])
    def gettimesheets_user():
        if 'org_code' not in request.args:
            return jsonify({'error': 'No Organization Code field provided.'}), 400

        if 'email_address' not in request.args:
            return jsonify({'error': 'No Email Address provided.'}), 400

        org_code = request.args['org_code']
        email_address = request.args['email_address'].strip()

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()  # DictCursor assumed

            sql = """
                SELECT
                    id,
                    email_address,
                    task,
                    start,
                    end,
                    hours,
                    created_datetime,
                    org_code,
                    task_type,
                    wo_code,
                    activities,
                    remarks
                FROM timesheets
                WHERE org_code = %s
                AND email_address = %s
                ORDER BY start DESC, id DESC
            """

            cur.execute(sql, (org_code, email_address))
            rows = cur.fetchall()  # dict rows

            return jsonify({'result': rows}), 200

        except Exception as e:
            print("gettimesheets_user error:", e)
            return jsonify({'error': str(e)}), 500


    # --------------------------------------------
    # Timesheet Engagement Matrix Report (MySQL)
    # --------------------------------------------
    @app.route('/timesheets_engagement_matrix', methods=['GET'])
    def timesheets_engagement_matrix():
        try:
            org_code = (request.args.get('org_code') or '').strip()
            year = (request.args.get('year') or '').strip()
            month = (request.args.get('month') or '').strip()  # optional: 1..12 or empty

            company = (request.args.get('company') or '').strip()          # business_units.code (optional)
            scope = (request.args.get('scope') or '').strip()              # services.service_id (optional)
            wo_search = (request.args.get('wo_search') or '').strip()
            employee_search = (request.args.get('employee_search') or '').strip()
            project_search = (request.args.get('project_search') or '').strip()

            if not org_code:
                return jsonify({"message": "org_code is required", "result": []}), 400

            if not year:
                year = str(datetime.now().year)

            y = int(year)

            # ✅ safer: filter by DATE boundaries (works for DATE or DATETIME columns)
            if month:
                m = int(month)
                last_day = calendar.monthrange(y, m)[1]
                start_date = f"{y:04d}-{m:02d}-01"
                end_date = f"{y:04d}-{m:02d}-{last_day:02d}"
            else:
                start_date = f"{y:04d}-01-01"
                end_date = f"{y:04d}-12-31"

            sql = """
                SELECT
                    IFNULL(s.description,'-') AS scope_desc,
                    IFNULL(bu.description,'-') AS company_desc,
                    IFNULL(w.project_name,'-') AS project_name,
                    TRIM(CONCAT(IFNULL(u.firstname,''), ' ', IFNULL(u.lastname,''))) AS engineer_name,
                    IFNULL(t.email_address,'') AS engineer_email,
                    IFNULL(t.task_type,'-') AS task_type,
                    MONTH(t.`start`) AS month_no,
                    SUM(IFNULL(t.hours,0)) AS hours
                FROM timesheets t
                LEFT JOIN work_orders w
                    ON t.wo_code = w.wo_code
                AND t.org_code = w.org_code

                -- ✅ prevent duplication if requested_services has multiple rows per wr_id
                LEFT JOIN (
                    SELECT org_code, wr_id, MIN(service_id) AS service_id
                    FROM requested_services
                    GROUP BY org_code, wr_id
                ) rs
                    ON w.wr_id = rs.wr_id
                AND w.org_code = rs.org_code

                LEFT JOIN services s
                    ON rs.service_id = s.service_id
                AND rs.org_code = s.org_code

                LEFT JOIN business_units bu
                    ON w.business_unit = bu.code
                AND w.org_code = bu.org_code

                LEFT JOIN app_users u
                    ON t.email_address = u.user
                AND t.org_code = u.org_code

                WHERE
                    t.org_code = %s
                    AND DATE(t.`start`) BETWEEN %s AND %s
            """

            params = [org_code, start_date, end_date]

            # ✅ optional filters
            if company:
                sql += " AND w.business_unit = %s "
                params.append(company)

            if scope:
                sql += " AND rs.service_id = %s "
                params.append(scope)

            if wo_search:
                sql += " AND t.wo_code LIKE %s "
                params.append(f"%{wo_search}%")

            if project_search:
                tokens = project_search.strip().split()
                boolean_query = " ".join([f"+{t}*" for t in tokens if t.strip()])
                sql += """
                    AND MATCH(w.project_name)
                        AGAINST (%s IN BOOLEAN MODE)
                """
                params.append(boolean_query)

            if employee_search:
                like = f"%{employee_search}%"
                sql += """
                    AND (
                        t.email_address LIKE %s
                        OR u.firstname LIKE %s
                        OR u.lastname LIKE %s
                        OR CONCAT(IFNULL(u.firstname,''), ' ', IFNULL(u.lastname,'')) LIKE %s
                    )
                """
                params.extend([like, like, like, like])

            sql += """
                GROUP BY
                    s.description, bu.description, w.project_name,
                    engineer_name, t.email_address,
                    t.task_type,
                    MONTH(t.`start`)
                ORDER BY
                    s.description, bu.description, w.project_name,
                    engineer_name, t.task_type, month_no
            """

            # ✅ execute (your project pattern)
            conn = dbconnect.getConnection()
            cur = conn.cursor()  # no dictionary=True
            cur.execute(sql, tuple(params))
            rows = cur.fetchall()
            cur.close()
            conn.close()

            # ✅ normalize output
            out = []
            for r in rows:
                if isinstance(r, dict):
                    scope_desc = (r.get('scope_desc') or '').strip() or '-'
                    company_desc = (r.get('company_desc') or '').strip() or '-'
                    project_name = (r.get('project_name') or '').strip() or '-'
                    engineer_name = (r.get('engineer_name') or '').strip()
                    engineer_email = (r.get('engineer_email') or '').strip() or '-'
                    task_type = (r.get('task_type') or '').strip() or '-'
                    month_no = int(r.get('month_no') or 0)
                    hours = float(r.get('hours') or 0)
                else:
                    # tuple/list (IMPORTANT: indexes shifted because task_type is added)
                    scope_desc = (r[0] or '').strip() or '-'
                    company_desc = (r[1] or '').strip() or '-'
                    project_name = (r[2] or '').strip() or '-'
                    engineer_name = (r[3] or '').strip()
                    engineer_email = (r[4] or '').strip() or '-'
                    task_type = (r[5] or '').strip() or '-'
                    month_no = int(r[6] or 0)
                    hours = float(r[7] or 0)

                if not engineer_name:
                    engineer_name = engineer_email

                out.append({
                    "scope_desc": scope_desc,
                    "company_desc": company_desc,
                    "project_name": project_name,
                    "engineer_name": engineer_name,
                    "engineer_email": engineer_email,
                    "task_type": task_type,
                    "month_no": month_no,
                    "hours": hours,
                })

            return jsonify({"message": "Engagement matrix retrieved successfully", "result": out}), 200

        except Exception as e:
            logger.exception("timesheets_engagement_matrix failed")
            return jsonify({"message": "Failed to retrieve engagement matrix", "error": str(e), "result": []}), 500


    #-- for searching work order via project name search
    @app.route('/work_orders_search', methods=['GET'])
    def work_orders_search():
        if 'org_code' not in request.args:
            return jsonify({'error': 'No Organization Code field provided.'}), 400

        if 'project_search' not in request.args:
            return jsonify({'error': 'No Project Search provided.'}), 400

        org_code = request.args['org_code']
        project_search = request.args['project_search'].strip()

        wo_search = request.args.get('wo_search', '').strip()
        status = request.args.get('status', '').strip()
        limit = request.args.get('limit', '20').strip()

        try:
            limit = int(limit)
            if limit <= 0:
                limit = 20
            if limit > 200:
                limit = 200
        except Exception:
            limit = 20

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()  # DictCursor assumed

            sql = """
                SELECT
                    wo_code,
                    project_name,
                    wo_description,
                    status,
                    location,
                    wo_number
                FROM work_orders
                WHERE org_code = %s
                AND project_name LIKE %s
            """
            params = [org_code, f"%{project_search}%"]

            # optional filters (matches your URL params)
            if wo_search:
                sql += " AND wo_code LIKE %s"
                params.append(f"%{wo_search}%")

            if status:
                sql += " AND status = %s"
                params.append(status)

            sql += """
                ORDER BY created_datetime DESC
                LIMIT %s
            """
            params.append(limit)

            cur.execute(sql, tuple(params))
            rows = cur.fetchall()

            return jsonify({'result': rows}), 200

        except Exception as e:
            print("work_orders_search error:", e)
            return jsonify({'error': str(e)}), 500

    
    #--- Getting work requests list for attaching to work order ---
    @app.route('/getworkrequests_for_attach', methods=['GET'])
    def getworkrequests_for_attach():
        try:
            org_code = (request.args.get('org_code') or '').strip()
            search = (request.args.get('search') or '').strip()
            businessunit_code = (request.args.get('businessunit_code') or '').strip()
            status = (request.args.get('status') or '').strip()  # optional exact match
            limit = request.args.get('limit', '30')

            if not org_code:
                return jsonify({"message": "org_code is required", "result": []}), 400

            # limit hygiene
            try:
                limit_i = int(limit)
            except:
                limit_i = 30
            if limit_i <= 0:
                limit_i = 30
            if limit_i > 100:
                limit_i = 100

            conn = dbconnect.getConnection()
            cur = conn.cursor()  # DictCursor assumed

            # ✅ Return ONLY 1 row per work_request, but keep your Flutter fields:
            #    - service_id (single value) : pick the latest requested_services.id per wr_id
            #    - scope (single value)      : based on that chosen service_id
            sql = """
                SELECT
                    wr.wr_id                              AS wr_id,
                    wr.wr_code                            AS wr_code,
                    wr.project_desc                       AS project_name,
                    bu.description                        AS business_unit,
                    wr.project_details                    AS details,
                    rs.service_id                         AS service_id,
                    svc.description                       AS scope,
                    wr.status                             AS status
                FROM work_requests wr
                LEFT JOIN business_units bu
                    ON wr.business_unit = bu.code

                LEFT JOIN (
                    SELECT x.wr_id, x.service_id
                    FROM requested_services x
                    INNER JOIN (
                        SELECT wr_id, MAX(id) AS max_id
                        FROM requested_services
                        WHERE org_code = %s
                        GROUP BY wr_id
                    ) y
                    ON x.wr_id = y.wr_id AND x.id = y.max_id
                ) rs
                    ON wr.wr_id = rs.wr_id

                LEFT JOIN services svc
                    ON rs.service_id = svc.service_id

                WHERE wr.status = 'Accepted' AND wr.org_code = %s
            """

            # params for subquery org_code + main org_code
            params = [org_code, org_code]

            # optional filters
            if businessunit_code:
                sql += " AND bu.description = %s "
                params.append(businessunit_code)

            if status:
                sql += " AND wr.status = %s "
                params.append(status)

            # 🔎 Search (case-insensitive for MySQL by using LOWER)
            if search:
                sql += """
                    AND (
                        LOWER(wr.wr_code) LIKE %s
                        OR LOWER(wr.project_desc) LIKE %s
                        OR LOWER(wr.project_details) LIKE %s
                    )
                """
                s = f"%{search.lower()}%"
                params.extend([s, s, s])

            # ✅ Order newest first
            sql += """
                ORDER BY wr.submitted_datetime DESC
                LIMIT %s
            """
            params.append(limit_i)

            cur.execute(sql, params)
            rows = cur.fetchall()

            return jsonify({"message": "success", "result": rows}), 200

        except Exception as e:
            return jsonify({"message": "error", "error": str(e), "result": []}), 500
        finally:
            try:
                cur.close()
                conn.close()
            except:
                pass


    #--- Get markup percent for a given org_code ---
    @app.route('/getmarkuppercent', methods=['GET'])
    def getmarkuppercent():
        org_code = (request.args.get('org_code') or '').strip()
        if not org_code:
            return jsonify({'error': 'No Organization Code field provided.'}), 400

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()  # DictCursor assumed

            cur.execute("""
                SELECT COALESCE(mark_up_percent, 0.00) AS mark_up_percent
                FROM sys_references
                WHERE org_code = %s
                LIMIT 1
            """, (org_code,))

            row = cur.fetchone()
            if not row:
                return jsonify({'org_code': org_code, 'mark_up_percent': 0.00}), 200

            return jsonify({
                'org_code': org_code,
                'mark_up_percent': float(row['mark_up_percent'])
            }), 200

        except Exception as e:
            return jsonify({'error': str(e)}), 500

        finally:
            try:
                cur.close()
                conn.close()
            except:
                pass

    # =========================================================
    # Get Work Order Revenue Batch Info
    # Used to validate WO Code and return WO description/status
    # =========================================================
    @app.route('/getworkorderrevenuebatchinfo', methods=['GET'])
    def getworkorderrevenuebatchinfo():

        if 'org_code' in request.args:
            org_code = request.args['org_code']
        else:
            return "Error: No Organization Code field provided. Please specify it."

        if 'wo_code' in request.args:
            wo_code = request.args['wo_code'].strip()
        else:
            return "Error: No Work Order Code field provided. Please specify it."

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            sql1 = """
                SELECT
                    a.wo_number AS wo_number,
                    a.wo_code AS wo_code,
                    a.wo_description AS wo_description,
                    a.status AS wo_status,
                    COALESCE(a.actual_total_cost, 0.00) AS actual_total_cost,
                    b.service_id AS service_id,
                    s.description AS service_description,
                    COALESCE(rac.code, '') AS revenue_code,
                    COALESCE(rac.account_title, '') AS revenue_account_title,

                    a.cost_type_used AS cost_type_used,
                    COALESCE(ce.total_cost, 0.00) AS estimate_high,
                    COALESCE(ce.total_cost_low, 0.00) AS estimate_low,
                    COALESCE(ce.total_cost_avg, 0.00) AS estimate_avg,
                    COALESCE(ce.mark_up_percent, 0.00) AS mark_up_percent

                FROM work_orders a

                LEFT JOIN (
                    SELECT
                        rs1.wr_id,
                        rs1.org_code,
                        rs1.service_id
                    FROM requested_services rs1
                    INNER JOIN (
                        SELECT
                            wr_id,
                            org_code,
                            MIN(id) AS first_id
                        FROM requested_services
                        GROUP BY wr_id, org_code
                    ) x
                        ON rs1.id = x.first_id
                        AND rs1.wr_id = x.wr_id
                        AND rs1.org_code = x.org_code
                ) b
                    ON a.wr_id = b.wr_id
                    AND a.org_code = b.org_code

                LEFT JOIN services s
                    ON b.service_id = s.service_id
                    AND b.org_code = s.org_code

                LEFT JOIN revenue_account_codes rac
                    ON rac.service_id = b.service_id
                    AND rac.org_code = a.org_code
                    AND rac.status = 1

                LEFT JOIN wo_cost_estimates ce
                    ON a.wo_number = ce.wo_number
                    AND a.org_code = ce.org_code

                WHERE UPPER(a.wo_code) = UPPER(%s)
                AND a.org_code = %s

                LIMIT 1
            """

            data1 = (wo_code, org_code)

            cur.execute(sql1, data1)
            result = cur.fetchone()

            app.config['JSONIFY_PRETTYPRINT_REGULAR'] = True

            if result:
                return jsonify({
                    "message": "Work order retrieved successfully",
                    "result": result
                }), 200
            else:
                return jsonify({
                    "message": "No work order found",
                    "result": None
                }), 404

        except Exception as e:
            return jsonify({
                "message": "Failed to retrieve work order",
                "error": str(e)
            }), 500

        finally:
            try:
                cur.close()
                conn.close()
            except:
                pass


    # =========================================================
    # Get Revenue Account Codes dropdown
    # Used by WO Revenue Entries dialog
    # =========================================================
    '''
    @app.route('/getrevenueaccountcodes', methods=['GET'])
    def getrevenueaccountcodes():
        org_code = request.args.get('org_code', '')

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            # ── Return ALL records (active AND inactive).
            # ── The Flutter client filters by status on its own.
            sql = """
                SELECT code, account_title, status, org_code
                FROM revenue_account_codes
                WHERE org_code = %s
                ORDER BY code ASC
            """
            cur.execute(sql, (org_code,))
            result = cur.fetchall()

            cur.close()
            conn.close()

            app.config['JSONIFY_PRETTYPRINT_REGULAR'] = True
            return jsonify({"result": result}), 200

        except Exception as e:
            print(str(e))
            return jsonify({"error": "Internal server error"}), 500
    '''
    @app.route('/getrevenueaccountcodes', methods=['GET'])
    def getrevenueaccountcodes():

        org_code = (request.args.get('org_code') or '').strip()

        if not org_code:
            return jsonify({
                "message": "No org_code field provided.",
                "result": []
            }), 400

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            sql = """
                SELECT
                    r.code,
                    r.account_title,
                    r.status,
                    r.org_code,
                    r.service_id,
                    COALESCE(s.description, '') AS scope
                FROM revenue_account_codes r
                LEFT JOIN services s
                    ON r.service_id = s.service_id
                AND r.org_code = s.org_code
                WHERE r.org_code = %s
                ORDER BY r.code ASC
            """

            cur.execute(sql, (org_code,))
            rows = cur.fetchall()

            result = []

            for row in rows:
                if isinstance(row, dict):
                    result.append({
                        "code": row.get("code"),
                        "account_title": row.get("account_title"),
                        "status": row.get("status"),
                        "org_code": row.get("org_code"),
                        "service_id": row.get("service_id"),
                        "scope": row.get("scope") or "",
                    })
                else:
                    result.append({
                        "code": row[0],
                        "account_title": row[1],
                        "status": row[2],
                        "org_code": row[3],
                        "service_id": row[4],
                        "scope": row[5] or "",
                    })

            cur.close()
            conn.close()

            return jsonify({
                "message": "Revenue account codes retrieved successfully",
                "result": result
            }), 200

        except Exception as e:
            print("getrevenueaccountcodes error:", str(e))
            return jsonify({
                "message": "Failed to retrieve revenue account codes",
                "error": str(e),
                "result": []
            }), 500
    

    ## ── Same fix applies to getserviceaccountcodes if it has the same issue ──

    @app.route('/getserviceaccountcodes', methods=['GET'])
    def getserviceaccountcodes():
        org_code = request.args.get('org_code', '')

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            # ── Return ALL records (active AND inactive).
            sql = """
                SELECT code, account_title, status, org_code, wac_category
                FROM service_account_codes
                WHERE org_code = %s
                ORDER BY code ASC
            """
            cur.execute(sql, (org_code,))
            result = cur.fetchall()

            cur.close()
            conn.close()

            app.config['JSONIFY_PRETTYPRINT_REGULAR'] = True
            return jsonify({"result": result}), 200

        except Exception as e:
            print(str(e))
            return jsonify({"error": "Internal server error"}), 500
        

    #--- get work order revenue entries ----#
    @app.route('/getworevenueentries', methods=['GET'])
    def getworevenueentries():
        if 'wo_number' not in request.args:
            return jsonify({
                "message": "No wo_number field provided. Please specify it.",
                "result": []
            }), 400

        if 'org_code' not in request.args:
            return jsonify({
                "message": "No org_code field provided. Please specify it.",
                "result": []
            }), 400

        wo_number = request.args['wo_number']
        org_code = request.args['org_code']

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            sql = """
                SELECT
                    r.revenue_id,
                    r.revenue_code,
                    r.amount,
                    a.id AS attachment_id,
                    a.file AS attachment_file
                FROM wo_revenues r
                LEFT JOIN revenue_attachments a
                    ON r.revenue_id = a.revenue_id
                    AND r.org_code = a.org_code
                WHERE r.wo_number = %s
                AND r.org_code = %s
                ORDER BY r.revenue_id ASC, a.id ASC
            """
            cur.execute(sql, (wo_number, org_code))
            rows = cur.fetchall()

            grouped = {}

            for row in rows:
                if isinstance(row, dict):
                    revenue_id = row.get("revenue_id")
                    revenue_code = row.get("revenue_code")
                    amount = float(row.get("amount") or 0)
                    attachment_id = row.get("attachment_id")
                    attachment_file = row.get("attachment_file")
                else:
                    revenue_id = row[0]
                    revenue_code = row[1]
                    amount = float(row[2] or 0)
                    attachment_id = row[3]
                    attachment_file = row[4]

                if revenue_id not in grouped:
                    grouped[revenue_id] = {
                        "revenue_id": revenue_id,
                        "revenue_code": revenue_code,
                        "amount": amount,
                        "attachments": []
                    }

                if attachment_id:
                    file_value = attachment_file or ''
                    file_name = file_value.split('/')[-1] if file_value else 'Attachment'

                    grouped[revenue_id]["attachments"].append({
                        "id": attachment_id,
                        "file_name": file_name,
                        "file_url": file_value,
                    })

            result = list(grouped.values())

            cur.close()
            conn.close()

            return jsonify({
                "message": "Success",
                "result": result
            }), 200

        except Exception as e:
            return jsonify({
                "message": str(e),
                "result": []
            }), 500


    #--- get revenue attachments by revenue_id ----#
    @app.route('/getrevenueattachments', methods=['GET'])
    def getrevenueattachments():
        if 'revenue_id' not in request.args:
            return jsonify({
                "message": "No revenue_id field provided. Please specify it.",
                "result": []
            }), 400

        if 'org_code' not in request.args:
            return jsonify({
                "message": "No org_code field provided. Please specify it.",
                "result": []
            }), 400

        revenue_id = request.args['revenue_id']
        org_code = request.args['org_code']

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            sql = """
                SELECT
                    id,
                    file
                FROM revenue_attachments
                WHERE revenue_id = %s
                AND org_code = %s
                ORDER BY id ASC
            """
            cur.execute(sql, (revenue_id, org_code))
            rows = cur.fetchall()

            result = []
            for row in rows:
                file_value = row[1] or ''
                file_name = file_value.split('/')[-1] if file_value else 'Attachment'

                result.append({
                    "id": row[0],
                    "file_name": file_name,
                    "file_url": file_value,
                })

            cur.close()
            conn.close()

            return jsonify({
                "message": "Revenue attachments retrieved successfully",
                "result": result
            }), 200

        except Exception as e:
            return jsonify({
                "message": str(e),
                "result": []
            }), 500


    #--- get work order cost of services entries ----#
    @app.route('/getwocostserviceentries', methods=['GET'])
    def getwocostserviceentries():

        if 'wo_number' not in request.args:
            return jsonify({
                "message": "No wo_number field provided. Please specify it.",
                "result": []
            }), 400

        if 'org_code' not in request.args:
            return jsonify({
                "message": "No org_code field provided. Please specify it.",
                "result": []
            }), 400

        wo_number = request.args['wo_number']
        org_code = request.args['org_code']

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            sql = """
                SELECT
                    c.cost_service_id,
                    c.service_code,
                    c.amount,
                    a.id AS attachment_id,
                    a.file AS attachment_file
                FROM wo_cost_services c
                LEFT JOIN cost_service_attachments a
                    ON c.cost_service_id = a.cost_service_id
                    AND c.org_code = a.org_code
                WHERE c.wo_number = %s
                AND c.org_code = %s
                ORDER BY c.cost_service_id ASC, a.id ASC
            """
            cur.execute(sql, (wo_number, org_code))
            rows = cur.fetchall()

            grouped = {}

            for row in rows:
                if isinstance(row, dict):
                    cost_service_id = row.get("cost_service_id")
                    service_code = row.get("service_code")
                    amount = float(row.get("amount") or 0)
                    attachment_id = row.get("attachment_id")
                    attachment_file = row.get("attachment_file")
                else:
                    cost_service_id = row[0]
                    service_code = row[1]
                    amount = float(row[2] or 0)
                    attachment_id = row[3]
                    attachment_file = row[4]

                if cost_service_id not in grouped:
                    grouped[cost_service_id] = {
                        "cost_service_id": cost_service_id,
                        "service_code": service_code,
                        "amount": amount,
                        "attachments": []
                    }

                if attachment_id:
                    file_value = attachment_file or ''
                    file_name = file_value.split('/')[-1] if file_value else 'Attachment'

                    grouped[cost_service_id]["attachments"].append({
                        "id": attachment_id,
                        "file_name": file_name,
                        "file_url": file_value,
                    })

            result = list(grouped.values())

            cur.close()
            conn.close()

            return jsonify({
                "message": "WO cost of services entries retrieved successfully",
                "result": result
            }), 200

        except Exception as e:
            print(str(e))
            return jsonify({
                "message": str(e),
                "result": []
            }), 500
    

    #--- get cost service attachments by cost_service_id ----#
    @app.route('/getcostserviceattachments', methods=['GET'])
    def getcostserviceattachments():
        if 'cost_service_id' not in request.args:
            return jsonify({
                "message": "No cost_service_id field provided. Please specify it.",
                "result": []
            }), 400

        if 'org_code' not in request.args:
            return jsonify({
                "message": "No org_code field provided. Please specify it.",
                "result": []
            }), 400

        cost_service_id = request.args['cost_service_id']
        org_code = request.args['org_code']

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            sql = """
                SELECT
                    id,
                    file
                FROM cost_service_attachments
                WHERE cost_service_id = %s
                AND org_code = %s
                ORDER BY id ASC
            """
            cur.execute(sql, (cost_service_id, org_code))
            rows = cur.fetchall()

            result = []
            for row in rows:
                file_value = row[1] or ''
                file_name = file_value.split('/')[-1] if file_value else 'Attachment'

                result.append({
                    "id": row[0],
                    "file_name": file_name,
                    "file_url": file_value,
                })

            cur.close()
            conn.close()

            return jsonify({
                "message": "Cost service attachments retrieved successfully",
                "result": result
            }), 200

        except Exception as e:
            return jsonify({
                "message": str(e),
                "result": []
            }), 500

    
    #--- get strict YTD income statement ----#
    @app.route('/getytdincomestatement', methods=['GET'])
    def getytdincomestatement():

        if 'org_code' in request.args:
            org_code = request.args['org_code']
        else:
            return jsonify({
                "message": "No org_code field provided. Please specify it.",
                "result": None
            }), 400

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            # --- REVENUE : strict YTD based on created_datetime ---
            sql1 = """
                SELECT
                    COALESCE(s.description, '-') AS description,
                    COALESCE(SUM(r.amount), 0.00) AS amount
                FROM wo_revenues r
                LEFT JOIN services s
                    ON r.service_id = s.service_id
                    AND r.org_code = s.org_code
                WHERE r.org_code = %s
                AND r.created_datetime IS NOT NULL
                AND r.created_datetime >= MAKEDATE(YEAR(CURDATE()), 1)
                AND r.created_datetime < DATE_ADD(CURDATE(), INTERVAL 1 DAY)
                GROUP BY COALESCE(s.description, '-')
                ORDER BY COALESCE(s.description, '-')
            """
            data1 = (org_code,)
            cur.execute(sql1, data1)
            revenue_rows = cur.fetchall()

            revenue_result = []
            revenue_total = 0.0

            for row in revenue_rows:
                if isinstance(row, dict):
                    desc = row.get("description") or "-"
                    amt = float(row.get("amount") or 0)
                else:
                    desc = row[0] if row[0] is not None else "-"
                    amt = float(row[1] or 0)

                revenue_result.append({
                    "description": desc,
                    "amount": amt
                })
                revenue_total += amt

            # --- COST OF SERVICES : strict YTD based on created_datetime ---
            sql2 = """
                SELECT
                    COALESCE(sac.account_title, '-') AS description,
                    COALESCE(SUM(c.amount), 0.00) AS amount
                FROM wo_cost_services c
                LEFT JOIN service_account_codes sac
                    ON c.service_code = sac.code
                    AND c.org_code = sac.org_code
                WHERE c.org_code = %s
                AND c.created_datetime IS NOT NULL
                AND c.created_datetime >= MAKEDATE(YEAR(CURDATE()), 1)
                AND c.created_datetime < DATE_ADD(CURDATE(), INTERVAL 1 DAY)
                GROUP BY COALESCE(sac.account_title, '-')
                ORDER BY COALESCE(sac.account_title, '-')
            """
            data2 = (org_code,)
            cur.execute(sql2, data2)
            cost_rows = cur.fetchall()

            cost_result = []
            cost_total = 0.0

            for row in cost_rows:
                if isinstance(row, dict):
                    desc = row.get("description") or "-"
                    amt = float(row.get("amount") or 0)
                else:
                    desc = row[0] if row[0] is not None else "-"
                    amt = float(row[1] or 0)

                cost_result.append({
                    "description": desc,
                    "amount": amt
                })
                cost_total += amt

            gross_profit = revenue_total - cost_total

            app.config['JSONIFY_PRETTYPRINT_REGULAR'] = True

            return jsonify({
                "message": "Strict YTD income statement retrieved successfully",
                "result": {
                    "revenue": revenue_result,
                    "revenue_total": revenue_total,
                    "cost_of_services": cost_result,
                    "cost_total": cost_total,
                    "gross_profit": gross_profit
                }
            }), 200

        except Exception as e:
            logger.info(str(e))
            return jsonify({
                "message": "Failed to retrieve strict YTD income statement",
                "result": None
            }), 500

        finally:
            try:
                cur.close()
                conn.close()
            except:
                pass


    #--- get work order budget vs actual variance ----#
    @app.route('/getwobudgetvsactualvariance', methods=['GET'])
    def getwobudgetvsactualvariance():

        if 'wo_number' in request.args:
            wo_number = request.args['wo_number']
        else:
            return jsonify({
                "message": "No wo_number field provided. Please specify it.",
                "result": None
            }), 400

        if 'org_code' in request.args:
            org_code = request.args['org_code']
        else:
            return jsonify({
                "message": "No org_code field provided. Please specify it.",
                "result": None
            }), 400

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            '''
            sql = """
                SELECT
                    x.description,
                    x.sort_order,
                    x.budget,
                    x.actual,
                    (x.budget - x.actual) AS variance
                FROM
                (
                    SELECT
                        'Materials' AS description,
                        1 AS sort_order,
                        COALESCE((
                            SELECT
                                CASE
                                    WHEN COALESCE(wce.mark_up_percent, 0) = 0 THEN COALESCE(wce.materials_cost, 0.00)
                                    ELSE COALESCE(wce.materials_cost, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                END
                            FROM wo_cost_estimates wce
                            WHERE wce.wo_number = %s
                            AND wce.org_code = %s
                            LIMIT 1
                        ), 0.00) AS budget,
                        COALESCE((
                            SELECT SUM(wcs.amount)
                            FROM wo_cost_services wcs
                            LEFT JOIN service_account_codes sac
                                ON sac.code = wcs.service_code
                                AND sac.org_code = wcs.org_code
                            WHERE wcs.wo_number = %s
                            AND wcs.org_code = %s
                            AND LOWER(TRIM(COALESCE(sac.wac_category, ''))) IN ('materials', 'material')
                        ), 0.00) AS actual

                    UNION ALL

                    SELECT
                        'Labor' AS description,
                        2 AS sort_order,
                        COALESCE((
                            SELECT
                                CASE
                                    WHEN COALESCE(wce.mark_up_percent, 0) = 0 THEN COALESCE(wce.labor_cost, 0.00)
                                    ELSE COALESCE(wce.labor_cost, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                END
                            FROM wo_cost_estimates wce
                            WHERE wce.wo_number = %s
                            AND wce.org_code = %s
                            LIMIT 1
                        ), 0.00) AS budget,
                        COALESCE((
                            SELECT SUM(wcs.amount)
                            FROM wo_cost_services wcs
                            LEFT JOIN service_account_codes sac
                                ON sac.code = wcs.service_code
                                AND sac.org_code = wcs.org_code
                            WHERE wcs.wo_number = %s
                            AND wcs.org_code = %s
                            AND LOWER(TRIM(COALESCE(sac.wac_category, ''))) IN ('labor', 'labour')
                        ), 0.00) AS actual

                    UNION ALL

                    SELECT
                        'Equipment' AS description,
                        3 AS sort_order,
                        COALESCE((
                            SELECT
                                CASE
                                    WHEN COALESCE(wce.mark_up_percent, 0) = 0 THEN COALESCE(wce.equipment_cost, 0.00)
                                    ELSE COALESCE(wce.equipment_cost, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                END
                            FROM wo_cost_estimates wce
                            WHERE wce.wo_number = %s
                            AND wce.org_code = %s
                            LIMIT 1
                        ), 0.00) AS budget,
                        COALESCE((
                            SELECT SUM(wcs.amount)
                            FROM wo_cost_services wcs
                            LEFT JOIN service_account_codes sac
                                ON sac.code = wcs.service_code
                                AND sac.org_code = wcs.org_code
                            WHERE wcs.wo_number = %s
                            AND wcs.org_code = %s
                            AND LOWER(TRIM(COALESCE(sac.wac_category, ''))) = 'equipment'
                        ), 0.00) AS actual

                    UNION ALL

                    SELECT
                        'Overhead' AS description,
                        4 AS sort_order,
                        COALESCE((
                            SELECT
                                CASE
                                    WHEN COALESCE(wce.mark_up_percent, 0) = 0 THEN COALESCE(wce.overhead_cost, 0.00)
                                    ELSE COALESCE(wce.overhead_cost, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                END
                            FROM wo_cost_estimates wce
                            WHERE wce.wo_number = %s
                            AND wce.org_code = %s
                            LIMIT 1
                        ), 0.00) AS budget,
                        COALESCE((
                            SELECT SUM(wcs.amount)
                            FROM wo_cost_services wcs
                            LEFT JOIN service_account_codes sac
                                ON sac.code = wcs.service_code
                                AND sac.org_code = wcs.org_code
                            WHERE wcs.wo_number = %s
                            AND wcs.org_code = %s
                            AND LOWER(TRIM(COALESCE(sac.wac_category, ''))) = 'overhead'
                        ), 0.00) AS actual

                    UNION ALL

                    SELECT
                        'Contingency' AS description,
                        5 AS sort_order,
                        COALESCE((
                            SELECT
                                CASE
                                    WHEN COALESCE(wce.mark_up_percent, 0) = 0 THEN COALESCE(wce.contingency_fund, 0.00)
                                    ELSE COALESCE(wce.contingency_fund, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                END
                            FROM wo_cost_estimates wce
                            WHERE wce.wo_number = %s
                            AND wce.org_code = %s
                            LIMIT 1
                        ), 0.00) AS budget,
                        COALESCE((
                            SELECT SUM(wcs.amount)
                            FROM wo_cost_services wcs
                            LEFT JOIN service_account_codes sac
                                ON sac.code = wcs.service_code
                                AND sac.org_code = wcs.org_code
                            WHERE wcs.wo_number = %s
                            AND wcs.org_code = %s
                            AND LOWER(TRIM(COALESCE(sac.wac_category, ''))) IN ('contingency', 'contingency fund')
                        ), 0.00) AS actual

                    UNION ALL

                    SELECT
                        'Discounts' AS description,
                        6 AS sort_order,
                        COALESCE((
                            SELECT
                                CASE
                                    WHEN COALESCE(wce.mark_up_percent, 0) = 0 THEN COALESCE(wce.discounts, 0.00)
                                    ELSE COALESCE(wce.discounts, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                END
                            FROM wo_cost_estimates wce
                            WHERE wce.wo_number = %s
                            AND wce.org_code = %s
                            LIMIT 1
                        ), 0.00) AS budget,
                        COALESCE((
                            SELECT SUM(wcs.amount)
                            FROM wo_cost_services wcs
                            LEFT JOIN service_account_codes sac
                                ON sac.code = wcs.service_code
                                AND sac.org_code = wcs.org_code
                            WHERE wcs.wo_number = %s
                            AND wcs.org_code = %s
                            AND LOWER(TRIM(COALESCE(sac.wac_category, ''))) IN ('discounts', 'discount')
                        ), 0.00) AS actual

                    UNION ALL

                    SELECT
                        'Total' AS description,
                        7 AS sort_order,
                        COALESCE((
                            SELECT
                                CASE
                                    WHEN COALESCE(wce.mark_up_percent, 0) = 0 THEN COALESCE(wce.total_cost, 0.00)
                                    ELSE COALESCE(wce.total_cost, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                END
                            FROM wo_cost_estimates wce
                            WHERE wce.wo_number = %s
                            AND wce.org_code = %s
                            LIMIT 1
                        ), 0.00) AS budget,
                        COALESCE((
                            SELECT SUM(wcs.amount)
                            FROM wo_cost_services wcs
                            WHERE wcs.wo_number = %s
                            AND wcs.org_code = %s
                        ), 0.00) AS actual
                ) x
                ORDER BY x.sort_order
            """
            '''
            sql = """
                SELECT
                    x.description,
                    x.sort_order,
                    x.budget,
                    x.actual,
                    (x.budget - x.actual) AS variance,
                    x.mark_up_percent
                FROM
                (
                    SELECT
                        'Materials' AS description,
                        1 AS sort_order,
                        COALESCE((
                            SELECT
                                CASE
                                    WHEN COALESCE(wce.mark_up_percent, 0) = 0 THEN
                                        CASE
                                            WHEN COALESCE(wo.cost_type_used, 'High') = 'Low' THEN COALESCE(wce.materials_cost_low, 0.00)
                                            WHEN COALESCE(wo.cost_type_used, 'High') = 'Average' THEN COALESCE(wce.materials_cost_avg, 0.00)
                                            ELSE COALESCE(wce.materials_cost, 0.00)
                                        END
                                    ELSE
                                        CASE
                                            WHEN COALESCE(wo.cost_type_used, 'High') = 'Low' THEN COALESCE(wce.materials_cost_low, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                            WHEN COALESCE(wo.cost_type_used, 'High') = 'Average' THEN COALESCE(wce.materials_cost_avg, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                            ELSE COALESCE(wce.materials_cost, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                        END
                                END
                            FROM wo_cost_estimates wce
                            LEFT JOIN work_orders wo
                                ON wo.wo_number = wce.wo_number
                                AND wo.org_code = wce.org_code
                            WHERE wce.wo_number = %s
                            AND wce.org_code = %s
                            LIMIT 1
                        ), 0.00) AS budget,
                        COALESCE((
                            SELECT SUM(wcs.amount)
                            FROM wo_cost_services wcs
                            LEFT JOIN service_account_codes sac
                                ON sac.code = wcs.service_code
                                AND sac.org_code = wcs.org_code
                            WHERE wcs.wo_number = %s
                            AND wcs.org_code = %s
                            AND LOWER(TRIM(COALESCE(sac.wac_category, ''))) IN ('materials', 'material')
                        ), 0.00) AS actual,
                        0.00 AS mark_up_percent

                    UNION ALL

                    SELECT
                        'Labor' AS description,
                        2 AS sort_order,
                        COALESCE((
                            SELECT
                                CASE
                                    WHEN COALESCE(wce.mark_up_percent, 0) = 0 THEN COALESCE(wce.labor_cost, 0.00)
                                    ELSE COALESCE(wce.labor_cost, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                END
                            FROM wo_cost_estimates wce
                            LEFT JOIN work_orders wo
                                ON wo.wo_number = wce.wo_number
                                AND wo.org_code = wce.org_code
                            WHERE wce.wo_number = %s
                            AND wce.org_code = %s
                            LIMIT 1
                        ), 0.00) AS budget,
                        COALESCE((
                            SELECT SUM(wcs.amount)
                            FROM wo_cost_services wcs
                            LEFT JOIN service_account_codes sac
                                ON sac.code = wcs.service_code
                                AND sac.org_code = wcs.org_code
                            WHERE wcs.wo_number = %s
                            AND wcs.org_code = %s
                            AND LOWER(TRIM(COALESCE(sac.wac_category, ''))) IN ('labor', 'labour')
                        ), 0.00) AS actual,
                        0.00 AS mark_up_percent

                    UNION ALL

                    SELECT
                        'Equipment' AS description,
                        3 AS sort_order,
                        COALESCE((
                            SELECT
                                CASE
                                    WHEN COALESCE(wce.mark_up_percent, 0) = 0 THEN COALESCE(wce.equipment_cost, 0.00)
                                    ELSE COALESCE(wce.equipment_cost, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                END
                            FROM wo_cost_estimates wce
                            LEFT JOIN work_orders wo
                                ON wo.wo_number = wce.wo_number
                                AND wo.org_code = wce.org_code
                            WHERE wce.wo_number = %s
                            AND wce.org_code = %s
                            LIMIT 1
                        ), 0.00) AS budget,
                        COALESCE((
                            SELECT SUM(wcs.amount)
                            FROM wo_cost_services wcs
                            LEFT JOIN service_account_codes sac
                                ON sac.code = wcs.service_code
                                AND sac.org_code = wcs.org_code
                            WHERE wcs.wo_number = %s
                            AND wcs.org_code = %s
                            AND LOWER(TRIM(COALESCE(sac.wac_category, ''))) = 'equipment'
                        ), 0.00) AS actual,
                        0.00 AS mark_up_percent

                    UNION ALL

                    SELECT
                        'Overhead' AS description,
                        4 AS sort_order,
                        COALESCE((
                            SELECT
                                CASE
                                    WHEN COALESCE(wce.mark_up_percent, 0) = 0 THEN COALESCE(wce.overhead_cost, 0.00)
                                    ELSE COALESCE(wce.overhead_cost, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                END
                            FROM wo_cost_estimates wce
                            LEFT JOIN work_orders wo
                                ON wo.wo_number = wce.wo_number
                                AND wo.org_code = wce.org_code
                            WHERE wce.wo_number = %s
                            AND wce.org_code = %s
                            LIMIT 1
                        ), 0.00) AS budget,
                        COALESCE((
                            SELECT SUM(wcs.amount)
                            FROM wo_cost_services wcs
                            LEFT JOIN service_account_codes sac
                                ON sac.code = wcs.service_code
                                AND sac.org_code = wcs.org_code
                            WHERE wcs.wo_number = %s
                            AND wcs.org_code = %s
                            AND LOWER(TRIM(COALESCE(sac.wac_category, ''))) = 'overhead'
                        ), 0.00) AS actual,
                        0.00 AS mark_up_percent

                    UNION ALL

                    SELECT
                        'Contingency' AS description,
                        5 AS sort_order,
                        COALESCE((
                            SELECT
                                CASE
                                    WHEN COALESCE(wce.mark_up_percent, 0) = 0 THEN COALESCE(wce.contingency_fund, 0.00)
                                    ELSE COALESCE(wce.contingency_fund, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                END
                            FROM wo_cost_estimates wce
                            LEFT JOIN work_orders wo
                                ON wo.wo_number = wce.wo_number
                                AND wo.org_code = wce.org_code
                            WHERE wce.wo_number = %s
                            AND wce.org_code = %s
                            LIMIT 1
                        ), 0.00) AS budget,
                        COALESCE((
                            SELECT SUM(wcs.amount)
                            FROM wo_cost_services wcs
                            LEFT JOIN service_account_codes sac
                                ON sac.code = wcs.service_code
                                AND sac.org_code = wcs.org_code
                            WHERE wcs.wo_number = %s
                            AND wcs.org_code = %s
                            AND LOWER(TRIM(COALESCE(sac.wac_category, ''))) IN ('contingency', 'contingency fund')
                        ), 0.00) AS actual,
                        0.00 AS mark_up_percent

                    UNION ALL

                    SELECT
                        'Discounts' AS description,
                        6 AS sort_order,
                        COALESCE((
                            SELECT
                                CASE
                                    WHEN COALESCE(wce.mark_up_percent, 0) = 0 THEN COALESCE(wce.discounts, 0.00)
                                    ELSE COALESCE(wce.discounts, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                END
                            FROM wo_cost_estimates wce
                            LEFT JOIN work_orders wo
                                ON wo.wo_number = wce.wo_number
                                AND wo.org_code = wce.org_code
                            WHERE wce.wo_number = %s
                            AND wce.org_code = %s
                            LIMIT 1
                        ), 0.00) AS budget,
                        COALESCE((
                            SELECT SUM(wcs.amount)
                            FROM wo_cost_services wcs
                            LEFT JOIN service_account_codes sac
                                ON sac.code = wcs.service_code
                                AND sac.org_code = wcs.org_code
                            WHERE wcs.wo_number = %s
                            AND wcs.org_code = %s
                            AND LOWER(TRIM(COALESCE(sac.wac_category, ''))) IN ('discounts', 'discount')
                        ), 0.00) AS actual,
                        0.00 AS mark_up_percent

                    UNION ALL

                    SELECT
                        'Total' AS description,
                        7 AS sort_order,
                        COALESCE((
                            SELECT
                                CASE
                                    WHEN COALESCE(wce.mark_up_percent, 0) = 0 THEN
                                        CASE
                                            WHEN COALESCE(wo.cost_type_used, 'High') = 'Low' THEN COALESCE(wce.total_cost_low, 0.00)
                                            WHEN COALESCE(wo.cost_type_used, 'High') = 'Average' THEN COALESCE(wce.total_cost_avg, 0.00)
                                            ELSE COALESCE(wce.total_cost, 0.00)
                                        END
                                    ELSE
                                        CASE
                                            WHEN COALESCE(wo.cost_type_used, 'High') = 'Low' THEN COALESCE(wce.total_cost_low, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                            WHEN COALESCE(wo.cost_type_used, 'High') = 'Average' THEN COALESCE(wce.total_cost_avg, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                            ELSE COALESCE(wce.total_cost, 0.00) / (1 + (COALESCE(wce.mark_up_percent, 0) / 100))
                                        END
                                END
                            FROM wo_cost_estimates wce
                            LEFT JOIN work_orders wo
                                ON wo.wo_number = wce.wo_number
                                AND wo.org_code = wce.org_code
                            WHERE wce.wo_number = %s
                            AND wce.org_code = %s
                            LIMIT 1
                        ), 0.00) AS budget,
                        COALESCE((
                            SELECT SUM(wcs.amount)
                            FROM wo_cost_services wcs
                            WHERE wcs.wo_number = %s
                            AND wcs.org_code = %s
                        ), 0.00) AS actual,
                        COALESCE((
                            SELECT COALESCE(wce.mark_up_percent, 0.00)
                            FROM wo_cost_estimates wce
                            LEFT JOIN work_orders wo
                                ON wo.wo_number = wce.wo_number
                                AND wo.org_code = wce.org_code
                            WHERE wce.wo_number = %s
                            AND wce.org_code = %s
                            LIMIT 1
                        ), 0.00) AS mark_up_percent
                ) x
                ORDER BY x.sort_order
            """

            '''
            data = (
                wo_number, org_code, wo_number, org_code,   # Materials
                wo_number, org_code, wo_number, org_code,   # Labor
                wo_number, org_code, wo_number, org_code,   # Equipment
                wo_number, org_code, wo_number, org_code,   # Overhead
                wo_number, org_code, wo_number, org_code,   # Contingency
                wo_number, org_code, wo_number, org_code,   # Discounts
                wo_number, org_code, wo_number, org_code    # Total
            )
            '''
            data = (
                wo_number, org_code, wo_number, org_code,   # Materials
                wo_number, org_code, wo_number, org_code,   # Labor
                wo_number, org_code, wo_number, org_code,   # Equipment
                wo_number, org_code, wo_number, org_code,   # Overhead
                wo_number, org_code, wo_number, org_code,   # Contingency
                wo_number, org_code, wo_number, org_code,   # Discounts
                wo_number, org_code, wo_number, org_code,   # Total budget + actual
                wo_number, org_code                         # Total mark_up_percent
            )

            cur.execute(sql, data)
            rows = cur.fetchall()

            result = []
            for row in rows:
                '''
                result.append({
                    "description": row["description"],
                    "budget": float(row["budget"] or 0),
                    "actual": float(row["actual"] or 0),
                    "variance": float(row["variance"] or 0),
                })
                '''
                result.append({
                    "description": row["description"],
                    "budget": float(row["budget"] or 0),
                    "actual": float(row["actual"] or 0),
                    "variance": float(row["variance"] or 0),
                    "mark_up_percent": float(row["mark_up_percent"] or 0),
                })

            cur.close()
            conn.close()

            return jsonify({
                "message": "Budget vs actual variance retrieved successfully.",
                "result": result
            }), 200

        except Exception as e:
            return jsonify({
                "message": repr(e),
                "result": []
            }), 500

    #==========================================================================
    
    #--- generate client proposal document for work order ----#
    @app.route('/generateproposal', methods=['GET'])
    def generateproposal():

        import os
        import time
        import tempfile
        import subprocess
        import traceback
        from datetime import datetime

        from flask import request, jsonify, send_file
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        from io import BytesIO
        from docx.shared import Inches

        wo_number = request.args.get('wo_number')
        org_code = request.args.get('org_code')
        format_type = (request.args.get('format') or 'docx').strip().lower()
        template_type = (request.args.get('template_type') or 'client').strip().lower()

        # Editable fields from frontend
        recipient_name_param = request.args.get('recipient_name')
        recipient_role_param = request.args.get('recipient_role')
        recipient_company_param = request.args.get('recipient_company')

        scope_of_work_param = request.args.get('scope_of_work')
        timeline_param = request.args.get('timeline')
        terms_conditions_param = request.args.get('terms_conditions')
        title_after_terms_param = request.args.get('title_after_terms')

        payment_terms_param = request.args.get('payment_terms')
        validity_days_param = request.args.get('validity_days')
        customer_name_param = request.args.get('customer_name')

        if not wo_number:
            return jsonify({"error": "1", "message": "No wo_number field provided."}), 400

        if not org_code:
            return jsonify({"error": "1", "message": "No org_code field provided."}), 400

        if format_type not in ['docx', 'pdf']:
            return jsonify({"error": "1", "message": "Invalid format. Use docx or pdf."}), 400

        conn = None
        cur = None

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            def getv(row, key, idx=None, default=''):
                if not row:
                    return default
                if isinstance(row, dict):
                    return row.get(key, default)
                try:
                    if idx is not None:
                        return row[idx]
                except Exception:
                    pass
                return default

            def text(v):
                if v is None:
                    return ''
                return str(v).strip()

            def money(v):
                try:
                    return f"{float(v or 0):,.2f}"
                except Exception:
                    return "0.00"

            def as_float(v):
                try:
                    return float(v or 0)
                except Exception:
                    return 0.0

            def safe_filename(v):
                return (
                    text(v)
                    .replace('/', '-')
                    .replace('\\', '-')
                    .replace(':', '-')
                    .replace('"', '')
                    .replace("'", '')
                )

            def set_cell_text(cell, value, bold=False, font_size=10):
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(text(value))
                run.bold = bold
                run.font.size = Pt(font_size)

            def set_cell_shading(cell, fill):
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), fill)
                tc_pr.append(shd)

            def center_cell(cell):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            def add_section_title(doc_obj, title):
                p = doc_obj.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                r = p.add_run(title)
                r.bold = True
                r.font.size = Pt(12)
                return p

            def convert_docx_to_pdf(docx_path):
                out_dir = os.path.dirname(docx_path)
                pdf_path = docx_path.replace(".docx", ".pdf")

                result = subprocess.run(
                    [
                        "libreoffice",
                        "--headless",
                        "--nologo",
                        "--nofirststartwizard",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        out_dir,
                        docx_path,
                    ],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                )

                if result.returncode != 0:
                    raise Exception(result.stderr or result.stdout or "LibreOffice conversion failed.")

                for _ in range(10):
                    if os.path.exists(pdf_path):
                        return pdf_path
                    time.sleep(0.5)

                raise Exception("PDF conversion completed but PDF file was not found.")

            # ----------------------------
            # 1. Work Order + WR info
            # ----------------------------
            sql1 = """
                SELECT
                    a.wo_number,
                    a.wo_code,
                    a.wo_description,
                    a.project_name,
                    a.project_description,
                    a.location,
                    a.business_unit,
                    bu.description AS business_unit_desc,
                    a.status AS wo_status,
                    a.proposal_status,
                    a.created_datetime,
                    a.due_date,
                    a.job_start_date,
                    a.job_end_date,
                    a.wr_id,
                    wr.firstname,
                    wr.middlename,
                    wr.lastname,
                    wr.email_address,
                    wr.project_location,
                    wr.project_desc,
                    wr.project_details,
                    wr.business_unit AS wr_business_unit,
                    wrbu.description AS wr_business_unit_desc
                FROM work_orders a
                LEFT JOIN business_units bu
                    ON a.business_unit = bu.code
                    AND a.org_code = bu.org_code
                LEFT JOIN work_requests wr
                    ON a.wr_id = wr.wr_id
                    AND a.org_code = wr.org_code
                LEFT JOIN business_units wrbu
                    ON wr.business_unit = wrbu.code
                    AND wr.org_code = wrbu.org_code
                WHERE a.wo_number = %s
                AND a.org_code = %s
                LIMIT 1
            """
            cur.execute(sql1, (wo_number, org_code))
            wo = cur.fetchone()

            if not wo:
                return jsonify({"error": "1", "message": "No work order found."}), 404

            wo_number_value = text(getv(wo, 'wo_number', 0, wo_number))
            wo_code = text(getv(wo, 'wo_code', 1, wo_number))
            wo_description = text(getv(wo, 'wo_description', 2, ''))
            project_name = text(getv(wo, 'project_name', 3, ''))
            project_description = text(getv(wo, 'project_description', 4, ''))
            location = text(getv(wo, 'location', 5, ''))
            created_datetime = getv(wo, 'created_datetime', 10, None)
            job_start_date = text(getv(wo, 'job_start_date', 12, ''))
            job_end_date = text(getv(wo, 'job_end_date', 13, ''))

            firstname = text(getv(wo, 'firstname', 15, ''))
            lastname = text(getv(wo, 'lastname', 17, ''))
            email_address = text(getv(wo, 'email_address', 18, ''))
            project_location = text(getv(wo, 'project_location', 19, ''))
            wr_description = text(getv(wo, 'project_desc', 20, ''))
            project_details = text(getv(wo, 'project_details', 21, ''))
            wr_business_unit_desc = text(getv(wo, 'wr_business_unit_desc', 23, ''))

            default_name = " ".join([firstname, lastname]).replace("  ", " ").strip()
            if not default_name:
                default_name = email_address

            recipient_name = text(recipient_name_param) or default_name
            recipient_role = text(recipient_role_param)
            recipient_company = text(recipient_company_param) or wr_business_unit_desc

            try:
                if created_datetime:
                    if isinstance(created_datetime, str):
                        proposal_date = datetime.strptime(created_datetime[:10], "%Y-%m-%d").strftime("%B %d, %Y")
                    else:
                        proposal_date = created_datetime.strftime("%B %d, %Y")
                else:
                    proposal_date = datetime.now().strftime("%B %d, %Y")
            except Exception:
                proposal_date = datetime.now().strftime("%B %d, %Y")

            proposal_title = project_name or wr_description or wo_description or "Proposal"
            wr_desc = wr_description or wo_description or proposal_title
            scope_of_work = text(scope_of_work_param) or project_description or project_details or wo_description or "-"

            timeline_text = text(timeline_param)
            if not timeline_text:
                if job_start_date or job_end_date:
                    timeline_text = f"{job_start_date or '-'} to {job_end_date or '-'}"
                else:
                    timeline_text = "Timeline shall be agreed upon by both parties."

            terms_conditions = text(terms_conditions_param)
            title_after_terms = text(title_after_terms_param)

            # ----------------------------
            # 2. Cost estimate
            # ----------------------------
            sql2 = """
                SELECT
                    COALESCE(e.materials_cost, 0.00) AS materials_cost,
                    COALESCE(e.labor_cost, 0.00) AS labor_cost,
                    COALESCE(e.equipment_cost, 0.00) AS equipment_cost,
                    COALESCE(e.overhead_cost, 0.00) AS overhead_cost,
                    COALESCE(e.contingency_fund, 0.00) AS contingency_fund,
                    COALESCE(e.discounts, 0.00) AS discounts,

                    COALESCE(e.total_cost, 0.00) AS total_cost_high,
                    COALESCE(e.total_cost_low, 0.00) AS total_cost_low,
                    COALESCE(e.total_cost_avg, 0.00) AS total_cost_avg,

                    COALESCE(e.tax, 0.00) AS tax,
                    COALESCE(e.payment_terms, '') AS payment_terms,
                    COALESCE(e.mark_up_percent, 0.00) AS mark_up_percent,
                    COALESCE(w.cost_type_used, 'High') AS cost_type_used
                FROM wo_cost_estimates e
                LEFT JOIN work_orders w
                    ON e.wo_number = w.wo_number
                    AND e.org_code = w.org_code
                WHERE e.wo_number = %s
                AND e.org_code = %s
                LIMIT 1
            """
            cur.execute(sql2, (wo_number, org_code))
            estimate = cur.fetchone()

            materials_cost = as_float(getv(estimate, 'materials_cost', 0, 0))
            labor_cost = as_float(getv(estimate, 'labor_cost', 1, 0))
            equipment_cost = as_float(getv(estimate, 'equipment_cost', 2, 0))
            overhead_cost = as_float(getv(estimate, 'overhead_cost', 3, 0))
            contingency_fund = as_float(getv(estimate, 'contingency_fund', 4, 0))
            discounts = as_float(getv(estimate, 'discounts', 5, 0))

            total_cost_high = as_float(getv(estimate, 'total_cost_high', 6, 0))
            total_cost_low = as_float(getv(estimate, 'total_cost_low', 7, 0))
            total_cost_avg = as_float(getv(estimate, 'total_cost_avg', 8, 0))

            tax = as_float(getv(estimate, 'tax', 9, 0))
            payment_terms = text(getv(estimate, 'payment_terms', 10, ''))
            mark_up_percent = as_float(getv(estimate, 'mark_up_percent', 11, 0))
            cost_type_used = text(getv(estimate, 'cost_type_used', 12, 'High')).lower()

            if cost_type_used == "low":
                base_total_cost = total_cost_low
            elif cost_type_used in ["average", "avg"]:
                base_total_cost = total_cost_avg
            else:
                base_total_cost = total_cost_high

            if base_total_cost <= 0:
                base_total_cost = (
                    materials_cost +
                    labor_cost +
                    equipment_cost +
                    overhead_cost +
                    contingency_fund -
                    discounts +
                    tax
                )

            if mark_up_percent > 0:
                total_cost = base_total_cost + (base_total_cost * (mark_up_percent / 100))
            else:
                total_cost = base_total_cost

            # ----------------------------
            # 3. Work Order Team Signatories
            # ----------------------------
            
            sql_team = """
                SELECT
                    wot.role,
                    au.firstname,
                    au.lastname,
                    au.position_title,
                    au.signature_file
                FROM work_order_team wot
                LEFT JOIN app_users au
                    ON wot.user = au.user
                    AND wot.org_code = au.org_code
                WHERE wot.wo_number = %s
                AND wot.org_code = %s
            """
            cur.execute(sql_team, (wo_number, org_code))
            team_rows = cur.fetchall()

            planner_name = ""
            planner_position = ""
            planner_signature = None

            team_lead_name = ""
            team_lead_position = ""
            team_lead_signature = None

            manager_name = ""
            manager_position = ""
            manager_signature = None

            for row in team_rows:
                role = text(getv(row, 'role', 0, '')).lower()
                fname = text(getv(row, 'firstname', 1, ''))
                lname = text(getv(row, 'lastname', 2, ''))
                position = text(getv(row, 'position_title', 3, ''))
                signature = getv(row, 'signature_file', 4, None)

                full_name = f"{fname} {lname}".strip()

                if role == "planner":
                    planner_name = full_name
                    planner_position = position
                    planner_signature = signature
                elif role == "team lead":
                    team_lead_name = full_name
                    team_lead_position = position
                    team_lead_signature = signature
                elif role == "manager":
                    manager_name = full_name
                    manager_position = position
                    manager_signature = signature

            # ----------------------------
            # 3.1 Approval Status Flags
            # ----------------------------
            sql_approval = """
                SELECT action_status
                FROM approval_requests
                WHERE txn_reference = %s
                AND org_code = %s
                AND action_status IN ('Team Lead Approved', 'Overall Lead Approved')
            """
            cur.execute(sql_approval, (wo_number, org_code))
            approval_rows = cur.fetchall()

            has_team_lead_approved = False
            has_overall_lead_approved = False

            for row in approval_rows:
                approval_status = text(getv(row, 'action_status', 0, '')).strip()

                if approval_status == "Team Lead Approved":
                    has_team_lead_approved = True

                if approval_status == "Overall Lead Approved":
                    has_overall_lead_approved = True
                    has_team_lead_approved = True

            # ----------------------------
            # 4. Build DOCX matching template
            # ----------------------------
            doc = Document()

            section = doc.sections[0]
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(10)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"{wo_code} | {wr_desc}")
            r.bold = True
            r.font.size = Pt(10)

            doc.add_paragraph("")

            # --- BIG SPACE FOR LETTERHEAD / LOGO ---
            '''
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(80)  # 👈 adjust this (60–120 depending on your layout)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = p.add_run(proposal_date)
            run.font.size = Pt(11)

            doc.add_paragraph("")
            '''
            # ---------------------------------
            # COMPANY LOGO + DATE
            # ---------------------------------

            logo_path = "/home/ubuntu/bloomerapi/static/v1e_logo.png"

            # Use a table for precise alignment
            header_tbl = doc.add_table(rows=2, cols=1)
            header_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Remove table borders
            tbl = header_tbl._tbl
            tblPr = tbl.tblPr
            tblBorders = OxmlElement('w:tblBorders')

            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tblBorders.append(border)

            tblPr.append(tblBorders)

            # -----------------------------
            # LOGO ROW
            # -----------------------------
            logo_cell = header_tbl.cell(0, 0)
            logo_cell.text = ''

            p_logo = logo_cell.paragraphs[0]
            p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_logo.paragraph_format.left_indent = Inches(-0.10)
            p_logo.paragraph_format.first_line_indent = Inches(0)
            p_logo.paragraph_format.space_before = Pt(0)
            p_logo.paragraph_format.space_after = Pt(0)

            #run_logo = p_logo.add_run()
            #run_logo.add_picture(logo_path, width=Inches(2.4))

            if os.path.exists(logo_path):
                run_logo = p_logo.add_run()

                run_logo.add_picture(
                    logo_path,
                    width=Inches(2.4)
                )

            # -----------------------------
            # DATE ROW
            # -----------------------------
            date_cell = header_tbl.cell(1, 0)
            date_cell.text = ''

            p_date = date_cell.paragraphs[0]
            p_date.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_date.paragraph_format.left_indent = Inches(0)
            p_date.paragraph_format.first_line_indent = Inches(0)
            p_date.paragraph_format.space_before = Pt(6)
            p_date.paragraph_format.space_after = Pt(0)

            run = p_date.add_run(proposal_date)
            run.font.size = Pt(11)

            doc.add_paragraph("")

            def add_compact_line(doc, value):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(value)
                run.font.size = Pt(10)
                return p

            add_compact_line(doc, recipient_name or "_____________________")
            add_compact_line(doc, recipient_role or "_____________________")
            add_compact_line(doc, recipient_company or "____________________________")
            doc.add_paragraph("")
            
            p = doc.add_paragraph()
            p.add_run("Subject: ").bold = True
            p.add_run(proposal_title)

            doc.add_paragraph("")
            doc.add_paragraph(f"Dear {recipient_name or 'Sir/Ma’am'};")
            doc.add_paragraph("")

            doc.add_paragraph(
                f"We are writing in response to your company’s request to submit a proposal for the {proposal_title}. "
                "Following this letter are the details of our offer for your evaluation."
            )

            doc.add_paragraph(
                "By signing this proposal, both parties agree to be bound by its terms and conditions. "
                "This proposal shall serve as a binding agreement upon acceptance and will supersede any prior "
                "agreements or understandings, whether written or verbal, related to the subject matter herein."
            )

            doc.add_paragraph(
                "Should you be amenable to the conditions stated herein, kindly signify your conformity at the "
                "space provided below and return to us one (1) signed copy of this proposal."
            )

            doc.add_paragraph("We look forward to your positive response and acceptance.")
            doc.add_paragraph("")
            doc.add_paragraph("Sincerely yours,")
            doc.add_paragraph("One Engineering – Vivant Infracore Holdings Inc.")
            doc.add_paragraph("")

            sig = doc.add_table(rows=4, cols=3)
            sig.alignment = WD_TABLE_ALIGNMENT.CENTER

            sign_headers = ["Prepared by:", "Endorsed by:", "Approved by:"]

            sign_names = [
                planner_name or "(WO Initiator)",
                team_lead_name or "(Team Lead Approver)",
                manager_name or "(Overall Team Lead Approver)",
            ]

            sign_positions = [
                planner_position or "Planner",
                team_lead_position or "Team Lead",
                manager_position or "Manager",
            ]

            signatures = [
                planner_signature,
                team_lead_signature if has_team_lead_approved else None,
                manager_signature if has_overall_lead_approved else None,
            ]

            for i in range(3):
                set_cell_text(sig.cell(0, i), sign_headers[i], bold=True)

                # Signature image / fallback
                sig_cell = sig.cell(1, i)
                sig_cell.text = ''
                p = sig_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

                if signatures[i]:
                    try:
                        image_stream = BytesIO(signatures[i])
                        run = p.add_run()
                        run.add_picture(image_stream, width=Inches(1.25))
                    except Exception:
                        run = p.add_run("")
                        run.font.size = Pt(10)
                else:
                    run = p.add_run("")
                    run.font.size = Pt(10)

                # Name - bold
                name_cell = sig.cell(2, i)
                name_cell.text = ''
                p = name_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(sign_names[i])
                run.bold = True
                run.font.size = Pt(10)

                # Position title
                pos_cell = sig.cell(3, i)
                pos_cell.text = ''
                p = pos_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(sign_positions[i])
                run.font.size = Pt(10)

                center_cell(sig.cell(0, i))
                center_cell(sig.cell(1, i))
                center_cell(sig.cell(2, i))
                center_cell(sig.cell(3, i))

            doc.add_paragraph("")

            doc.add_page_break()

            add_section_title(doc, "Scope of Work")
            doc.add_paragraph(scope_of_work)

            add_section_title(doc, "Timeline")
            doc.add_paragraph(timeline_text)

            add_section_title(doc, "Terms and Conditions")
            if terms_conditions:
                doc.add_paragraph(terms_conditions)

            if title_after_terms:
                add_section_title(doc, title_after_terms)

            add_section_title(doc, "Commercial Offer")

            doc.add_paragraph(
                f"For the aforementioned scope of work, the total contract price is PHP {money(total_cost)} only."
            )

            doc.add_paragraph(
                "The price excludes Value Added Tax (VAT) and any other applicable government charges. "
                "Any additional costs such as permits, clearances, etc., if required, shall be shouldered by the client."
            )

            add_section_title(doc, "Payment Terms")

            payment_terms_text = text(payment_terms_param)

            if payment_terms_text:
                for line in payment_terms_text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(f"• {line.strip()}")
            else:
                doc.add_paragraph("• Down Payment: 50% upon issuance of Purchase Order")
                doc.add_paragraph("• Final Payment: 50% upon issuance of Certificate of Completion and Technical Report")

            add_section_title(doc, "Validity of Offer")

            validity_days = text(validity_days_param) or "30"

            doc.add_paragraph(
                f"This proposal is valid for {validity_days} days from the date of issuance. "
                "Any changes beyond this period may be subject to price adjustments."
            )

            doc.add_paragraph("")
            doc.add_paragraph("")
            customer_name = text(customer_name_param)

            # Label
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            p.add_run("Customer’s Representative:")

            # Add more vertical space BEFORE name
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(30)  # 👈 bigger gap here
            p.paragraph_format.space_after = Pt(0)

            if customer_name:
                run = p.add_run(customer_name)
            else:
                run = p.add_run("__________________________________")

            run.font.size = Pt(10)

            # Printed Name (small font, tight spacing)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)

            run = p.add_run("(Printed Name & Signature)")
            run.font.size = Pt(9)  # 👈 smaller font

            file_name = f"Proposal_{safe_filename(wo_code or wo_number_value)}.docx"
            docx_path = os.path.join(tempfile.gettempdir(), file_name)

            doc.save(docx_path)

            if format_type == "pdf":
                pdf_path = convert_docx_to_pdf(docx_path)

                return send_file(
                    pdf_path,
                    as_attachment=True,
                    download_name=file_name.replace(".docx", ".pdf"),
                    mimetype="application/pdf",
                )

            return send_file(
                docx_path,
                as_attachment=True,
                download_name=file_name,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        except Exception as e:
            return jsonify({
                "error": "1",
                "message": "Failed to generate proposal",
                "details": str(e),
                "trace": traceback.format_exc(),
            }), 500

        finally:
            try:
                if cur:
                    cur.close()
            except Exception:
                pass

            try:
                if conn:
                    conn.close()
            except Exception:
                pass

#===========================================================================
    #--- get proposal preview data for work order ----#
    @app.route('/getproposaldata', methods=['GET'])
    def getproposaldata():

        if 'wo_number' in request.args:
            wo_number = request.args['wo_number']
        else:
            return jsonify({
                "message": "No wo_number field provided. Please specify it.",
                "result": None
            }), 400

        if 'org_code' in request.args:
            org_code = request.args['org_code']
        else:
            return jsonify({
                "message": "No org_code field provided. Please specify it.",
                "result": None
            }), 400

        template_type = (request.args.get('template_type') or 'client').strip().lower()

        conn = None
        cur = None

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            def getv(row, key, idx=None, default=''):
                if not row:
                    return default
                if isinstance(row, dict):
                    return row.get(key, default)
                try:
                    if idx is not None:
                        return row[idx]
                except:
                    pass
                return default

            def text(v):
                if v is None:
                    return ''
                return str(v).strip()

            def as_float(v):
                try:
                    return float(v or 0)
                except:
                    return 0.0

            # ---------------------------------------------------------
            # 1. WORK ORDER HEADER + CLIENT INFO
            # ---------------------------------------------------------
            sql1 = """
                SELECT
                    a.wo_number,
                    a.wo_code,
                    a.wo_description,
                    a.project_name,
                    a.project_description,
                    a.location,
                    a.business_unit,
                    bu.description AS business_unit_desc,
                    a.status AS wo_status,
                    a.proposal_status,
                    a.created_datetime,
                    a.due_date,
                    a.job_start_date,
                    a.job_end_date,
                    a.wr_id,
                    wr.firstname,
                    wr.middlename,
                    wr.lastname,
                    wr.email_address,
                    wr.project_location,
                    wr.project_desc,
                    wr.project_details,
                    wr.business_unit AS wr_business_unit,
                    wrbu.description AS wr_business_unit_desc
                FROM work_orders a
                LEFT JOIN business_units bu
                    ON a.business_unit = bu.code
                    AND a.org_code = bu.org_code
                LEFT JOIN work_requests wr
                    ON a.wr_id = wr.wr_id
                    AND a.org_code = wr.org_code
                LEFT JOIN business_units wrbu
                    ON wr.business_unit = wrbu.code
                    AND wr.org_code = wrbu.org_code
                WHERE a.wo_number = %s
                AND a.org_code = %s
                LIMIT 1
            """
            cur.execute(sql1, (wo_number, org_code))
            wo = cur.fetchone()

            if not wo:
                return jsonify({
                    "message": "No work order found.",
                    "result": None
                }), 404

            # ---------------------------------------------------------
            # 2. COST ESTIMATE SUMMARY
            # ---------------------------------------------------------
            sql2 = """
                SELECT
                    COALESCE(materials_cost, 0.00) AS materials_cost,
                    COALESCE(materials_cost_low, 0.00) AS materials_cost_low,
                    COALESCE(materials_cost_avg, 0.00) AS materials_cost_avg,
                    COALESCE(labor_cost, 0.00) AS labor_cost,
                    COALESCE(equipment_cost, 0.00) AS equipment_cost,
                    COALESCE(overhead_cost, 0.00) AS overhead_cost,
                    COALESCE(contingency_fund, 0.00) AS contingency_fund,
                    COALESCE(discounts, 0.00) AS discounts,
                    COALESCE(total_cost, 0.00) AS total_cost,
                    COALESCE(total_cost_low, 0.00) AS total_cost_low,
                    COALESCE(total_cost_avg, 0.00) AS total_cost_avg,
                    COALESCE(tax, 0.00) AS tax,
                    COALESCE(payment_terms, '') AS payment_terms,
                    COALESCE(mark_up_percent, 0.00) AS mark_up_percent
                FROM wo_cost_estimates
                WHERE wo_number = %s
                AND org_code = %s
                LIMIT 1
            """
            cur.execute(sql2, (wo_number, org_code))
            estimate = cur.fetchone()

            if not estimate:
                estimate = {
                    "materials_cost": 0.0,
                    "materials_cost_low": 0.0,
                    "materials_cost_avg": 0.0,
                    "labor_cost": 0.0,
                    "equipment_cost": 0.0,
                    "overhead_cost": 0.0,
                    "contingency_fund": 0.0,
                    "discounts": 0.0,
                    "total_cost": 0.0,
                    "total_cost_low": 0.0,
                    "total_cost_avg": 0.0,
                    "tax": 0.0,
                    "payment_terms": "",
                    "mark_up_percent": 0.0
                }

            # ---------------------------------------------------------
            # 3. BILL OF MATERIALS
            # ---------------------------------------------------------
            sql3 = """
                SELECT
                    x.task_number,
                    x.item_code,
                    COALESCE(pi.description, CONCAT('Item Code ', x.item_code)) AS item_description,
                    x.cu_code,
                    x.uom,
                    x.quantity,
                    x.unit_cost,
                    x.total_cost,
                    x.source_type
                FROM
                (
                    SELECT
                        task_number,
                        item_code,
                        cu_code,
                        uom,
                        quantity,
                        unit_cost,
                        total_cost,
                        'CU Item' AS source_type
                    FROM wo_task_physical_items
                    WHERE wo_number = %s
                    AND org_code = %s

                    UNION ALL

                    SELECT
                        task_number,
                        item_code,
                        NULL AS cu_code,
                        uom,
                        quantity,
                        unit_cost,
                        total_cost,
                        'Custom Item' AS source_type
                    FROM wo_task_physical_custom_items
                    WHERE wo_number = %s
                    AND org_code = %s
                ) x
                LEFT JOIN physical_items pi
                    ON x.item_code = pi.item_code
                    AND pi.org_code = %s
                ORDER BY x.task_number, x.source_type, item_description
            """
            cur.execute(sql3, (wo_number, org_code, wo_number, org_code, org_code))
            material_rows = cur.fetchall()

            # ---------------------------------------------------------
            # 4. LABOR BREAKDOWN
            # ---------------------------------------------------------
            sql4 = """
                SELECT
                    x.task_number,
                    x.item_code,
                    COALESCE(hi.description, CONCAT('Item Code ', x.item_code)) AS item_description,
                    x.cu_code,
                    x.uom,
                    x.quantity,
                    x.duration,
                    x.unit_cost,
                    x.total_cost,
                    x.source_type
                FROM
                (
                    SELECT
                        task_number,
                        item_code,
                        cu_code,
                        uom,
                        quantity,
                        duration,
                        unit_cost,
                        total_cost,
                        'CU Labor' AS source_type
                    FROM wo_task_human_items
                    WHERE wo_number = %s
                    AND org_code = %s

                    UNION ALL

                    SELECT
                        task_number,
                        item_code,
                        NULL AS cu_code,
                        uom,
                        quantity,
                        duration,
                        unit_cost,
                        total_cost,
                        'Custom Labor' AS source_type
                    FROM wo_task_human_custom_items
                    WHERE wo_number = %s
                    AND org_code = %s
                ) x
                LEFT JOIN human_items hi
                    ON x.item_code = hi.item_code
                    AND hi.org_code = %s
                ORDER BY x.task_number, x.source_type, item_description
            """
            cur.execute(sql4, (wo_number, org_code, wo_number, org_code, org_code))
            labor_rows = cur.fetchall()

            # ---------------------------------------------------------
            # 5. EQUIPMENT BREAKDOWN
            # ---------------------------------------------------------
            sql5 = """
                SELECT
                    x.task_number,
                    x.item_code,
                    COALESCE(ei.description, CONCAT('Item Code ', x.item_code)) AS item_description,
                    x.cu_code,
                    x.uom,
                    x.quantity,
                    x.equip_usage,
                    x.unit_cost,
                    x.total_cost,
                    x.source_type
                FROM
                (
                    SELECT
                        task_number,
                        item_code,
                        cu_code,
                        uom,
                        quantity,
                        equip_usage,
                        unit_cost,
                        total_cost,
                        'CU Equipment' AS source_type
                    FROM wo_task_physical_equip_items
                    WHERE wo_number = %s
                    AND org_code = %s

                    UNION ALL

                    SELECT
                        task_number,
                        item_code,
                        NULL AS cu_code,
                        uom,
                        quantity,
                        equip_usage,
                        unit_cost,
                        total_cost,
                        'Custom Equipment' AS source_type
                    FROM wo_task_physical_equip_custom_items
                    WHERE wo_number = %s
                    AND org_code = %s
                ) x
                LEFT JOIN physical_equip_items ei
                    ON x.item_code = ei.item_code
                    AND ei.org_code = %s
                ORDER BY x.task_number, x.source_type, item_description
            """
            cur.execute(sql5, (wo_number, org_code, wo_number, org_code, org_code))
            equipment_rows = cur.fetchall()

            materials_cost = as_float(getv(estimate, 'materials_cost', 0, 0))
            labor_cost = as_float(getv(estimate, 'labor_cost', 3, 0))
            equipment_cost = as_float(getv(estimate, 'equipment_cost', 4, 0))
            overhead_cost = as_float(getv(estimate, 'overhead_cost', 5, 0))
            contingency_fund = as_float(getv(estimate, 'contingency_fund', 6, 0))
            discounts = as_float(getv(estimate, 'discounts', 7, 0))
            total_cost = as_float(getv(estimate, 'total_cost', 8, 0))
            tax = as_float(getv(estimate, 'tax', 11, 0))
            payment_terms = text(getv(estimate, 'payment_terms', 12, ''))
            mark_up_percent = as_float(getv(estimate, 'mark_up_percent', 13, 0))

            subtotal_before_markup = (
                materials_cost +
                labor_cost +
                equipment_cost +
                overhead_cost +
                contingency_fund -
                discounts +
                tax
            )

            final_total = subtotal_before_markup if template_type == 'internal' else (total_cost if total_cost > 0 else subtotal_before_markup)

            client_name = " ".join([
                text(getv(wo, 'firstname', 15, '')),
                text(getv(wo, 'middlename', 16, '')),
                text(getv(wo, 'lastname', 17, ''))
            ]).replace("  ", " ").strip()

            if not client_name:
                client_name = text(getv(wo, 'email_address', 18, ''))

            result = {
                "header": {
                    "proposal_no": f"PRO-{text(getv(wo, 'wo_code', 1, wo_number))}",
                    "wo_number": text(getv(wo, 'wo_number', 0, wo_number)),
                    "wo_code": text(getv(wo, 'wo_code', 1, '')),
                    "project_name": (
                        text(getv(wo, 'project_name', 3, '')) or
                        text(getv(wo, 'project_desc', 20, '')) or
                        text(getv(wo, 'wo_description', 2, ''))
                    ),
                    "project_location": (
                        text(getv(wo, 'location', 5, '')) or
                        text(getv(wo, 'project_location', 19, ''))
                    ),
                    "client_name": " ".join([
                        text(getv(wo, 'firstname', 15, '')),
                        text(getv(wo, 'lastname', 17, ''))
                    ]).replace("  ", " ").strip() or text(getv(wo, 'email_address', 18, '')),
                    "recipient_role": "",
                    "company": text(getv(wo, 'wr_business_unit_desc', 23, '')),
                    "business_unit": text(getv(wo, 'business_unit_desc', 7, '')),
                    "wo_status": text(getv(wo, 'wo_status', 8, '')),
                    "proposal_status": text(getv(wo, 'proposal_status', 9, '')),
                    "proposal_date": text(getv(wo, 'created_datetime', 10, '')),
                },
                "project_description": (
                    text(getv(wo, 'wo_description', 2, '')) or
                    text(getv(wo, 'project_description', 4, '')) or
                    text(getv(wo, 'project_details', 21, ''))
                ),
                "cost_summary": {
                    "materials_cost": materials_cost,
                    "labor_cost": labor_cost,
                    "equipment_cost": equipment_cost,
                    "overhead_cost": overhead_cost,
                    "contingency_fund": contingency_fund,
                    "discounts": discounts,
                    "tax": tax,
                    "mark_up_percent": mark_up_percent,
                    "payment_terms": payment_terms,
                    "total_project_cost": final_total,
                },
                "materials": [
                    {
                        "task_number": getv(r, 'task_number', 0, ''),
                        "item_description": getv(r, 'item_description', 2, ''),
                        "uom": getv(r, 'uom', 4, ''),
                        "quantity": as_float(getv(r, 'quantity', 5, 0)),
                        "unit_cost": as_float(getv(r, 'unit_cost', 6, 0)),
                        "total_cost": as_float(getv(r, 'total_cost', 7, 0)),
                        "source_type": getv(r, 'source_type', 8, ''),
                    } for r in material_rows
                ],
                "labor": [
                    {
                        "task_number": getv(r, 'task_number', 0, ''),
                        "item_description": getv(r, 'item_description', 2, ''),
                        "uom": getv(r, 'uom', 4, ''),
                        "quantity": as_float(getv(r, 'quantity', 5, 0)),
                        "duration": as_float(getv(r, 'duration', 6, 0)),
                        "unit_cost": as_float(getv(r, 'unit_cost', 7, 0)),
                        "total_cost": as_float(getv(r, 'total_cost', 8, 0)),
                        "source_type": getv(r, 'source_type', 9, ''),
                    } for r in labor_rows
                ],
                "equipment": [
                    {
                        "task_number": getv(r, 'task_number', 0, ''),
                        "item_description": getv(r, 'item_description', 2, ''),
                        "uom": getv(r, 'uom', 4, ''),
                        "quantity": as_float(getv(r, 'quantity', 5, 0)),
                        "equip_usage": as_float(getv(r, 'equip_usage', 6, 0)),
                        "unit_cost": as_float(getv(r, 'unit_cost', 7, 0)),
                        "total_cost": as_float(getv(r, 'total_cost', 8, 0)),
                        "source_type": getv(r, 'source_type', 9, ''),
                    } for r in equipment_rows
                ],
                "terms_and_conditions": [
                    "This proposal is based on the approved scope and estimate generated from WACS.",
                    "Any additional work outside the approved scope shall be subject to separate evaluation and approval.",
                    "Prices are subject to validation based on final site conditions and client requirements.",
                    payment_terms if payment_terms else "Payment terms shall be discussed separately.",
                    "This document is system-generated from the Work and Cost Management System."
                ]
            }

            return jsonify({
                "message": "Proposal data retrieved successfully",
                "result": result
            }), 200

        except Exception as e:
            return jsonify({
                "message": "Failed to retrieve proposal data",
                "error": str(e),
                "result": None
            }), 500

        finally:
            try:
                if cur:
                    cur.close()
            except:
                pass
            try:
                if conn:
                    conn.close()
            except:
                pass

    
    #--- get proposal details for work order (for 2-proposal template) ----#
    @app.route('/getwoproposaldetails', methods=['GET'])
    def getwoproposaldetails():

        wo_number = request.args.get('wo_number')
        org_code = request.args.get('org_code')

        if not wo_number:
            return jsonify({"error": "1", "message": "No wo_number field provided.", "result": None}), 400

        if not org_code:
            return jsonify({"error": "1", "message": "No org_code field provided.", "result": None}), 400

        conn = None
        cur = None

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            sql = """
                SELECT
                    wo_number,
                    org_code,
                    recipient_name,
                    recipient_role,
                    recipient_company,
                    scope_of_work,
                    timeline,
                    terms_conditions,
                    title_after_terms,
                    payment_terms,
                    validity_days,
                    customer_name
                FROM wo_proposal_details
                WHERE wo_number = %s
                AND org_code = %s
                LIMIT 1
            """
            cur.execute(sql, (wo_number, org_code))
            row = cur.fetchone()

            def getv(row, key, idx=None, default=''):
                if not row:
                    return default
                if isinstance(row, dict):
                    return row.get(key, default)
                try:
                    if idx is not None:
                        return row[idx]
                except:
                    pass
                return default

            result = None

            if row:
                result = {
                    "wo_number": getv(row, "wo_number", 0, ""),
                    "org_code": getv(row, "org_code", 1, ""),
                    "recipient_name": getv(row, "recipient_name", 2, ""),
                    "recipient_role": getv(row, "recipient_role", 3, ""),
                    "recipient_company": getv(row, "recipient_company", 4, ""),
                    "scope_of_work": getv(row, "scope_of_work", 5, ""),
                    "timeline": getv(row, "timeline", 6, ""),
                    "terms_conditions": getv(row, "terms_conditions", 7, ""),
                    "title_after_terms": getv(row, "title_after_terms", 8, ""),
                    "payment_terms": getv(row, "payment_terms", 9, ""),
                    "validity_days": getv(row, "validity_days", 10, ""),
                    "customer_name": getv(row, "customer_name", 11, ""),
                }

            return jsonify({
                "error": "0",
                "message": "Proposal details retrieved successfully",
                "result": result
            }), 200

        except Exception as e:
            return jsonify({
                "error": "1",
                "message": "Failed to retrieve proposal details",
                "details": str(e),
                "result": None
            }), 500

        finally:
            try:
                if cur:
                    cur.close()
                if conn:
                    conn.close()
            except:
                pass

    
    #--- get user signature (for 2-proposal template) ----#
    @app.route('/getusersignature', methods=['GET'])
    def getusersignature():
        import base64

        org_code = request.args.get('org_code')
        user_id = request.args.get('id')

        if not org_code or not user_id:
            return jsonify({
                "error": "1",
                "message": "Missing org_code or id",
                "result": None
            }), 400

        conn = None
        cur = None

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            cur.execute("""
                SELECT signature_file
                FROM app_users
                WHERE id = %s
                AND org_code = %s
                LIMIT 1
            """, (user_id, org_code))

            row = cur.fetchone()

            if not row:
                return jsonify({
                    "error": "1",
                    "message": "User not found",
                    "result": None
                }), 404

            signature_file = row.get('signature_file') if isinstance(row, dict) else row[0]

            if not signature_file:
                return jsonify({
                    "error": "0",
                    "message": "No signature found",
                    "result": None
                }), 200

            signature_base64 = base64.b64encode(signature_file).decode('utf-8')

            return jsonify({
                "error": "0",
                "message": "Signature retrieved",
                "result": {
                    "signature_file_base64": f"data:image/png;base64,{signature_base64}"
                }
            }), 200

        except Exception as e:
            return jsonify({
                "error": "1",
                "message": "Failed to retrieve signature",
                "details": str(e),
                "result": None
            }), 500

        finally:
            try:
                if cur:
                    cur.close()
                if conn:
                    conn.close()
            except:
                pass

    #--- get services list (for revenue code & scope mapping) ----#
    @app.route('/getserviceslist2', methods=['GET'])
    def getserviceslist2():

        org_code = (request.args.get('org_code') or '').strip()

        if not org_code:
            return jsonify({
                "message": "No org_code field provided.",
                "result": []
            }), 400

        try:
            conn = dbconnect.getConnection()
            cur = conn.cursor()

            sql = """
                SELECT
                    service_id,
                    description
                FROM services
                WHERE status = 1
                AND org_code = %s
                ORDER BY sequence ASC, description ASC
            """

            cur.execute(sql, (org_code,))
            rows = cur.fetchall()

            result = []

            for row in rows:
                if isinstance(row, dict):
                    result.append({
                        "service_id": row.get("service_id"),
                        "description": row.get("description"),
                    })
                else:
                    result.append({
                        "service_id": row[0],
                        "description": row[1],
                    })

            cur.close()
            conn.close()

            return jsonify({
                "message": "Services retrieved successfully",
                "result": result
            }), 200

        except Exception as e:
            print("getserviceslist2 error:", str(e))
            return jsonify({
                "message": "Failed to retrieve services",
                "error": str(e),
                "result": []
            }), 500